from datetime import datetime, timedelta
import json
import io
import os
import tempfile
from unittest import mock

import docx
from django.conf import settings
from django.contrib.auth.models import User
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils.timezone import now as django_now

from .models import CV, Job, SearchRun, UserPreferences
from .services.matching import (
    detect_sponsorship,
    parse_salary,
    salary_meets_threshold,
    _parse_match_response,
)
from .services.pdf_generator import build_pdf_filename, generate_tailored_pdf
from .services.tailoring import tailor_cv_for_job

# Isolate uploaded/generated media into a temp dir for the whole test module.
_TEST_MEDIA = tempfile.mkdtemp(prefix='ja_test_media_')


def build_docx_bytes():
    document = docx.Document()
    document.add_paragraph('Jane Doe - Software Engineer')
    document.add_paragraph('Skills: Python, Django, SQL')
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


class PreferencesViewTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username='carol', password='pw12345!')

    def test_dashboard_requires_login(self):
        response = self.client.get(reverse('dashboard'))
        self.assertEqual(response.status_code, 302)
        self.assertIn(reverse('login'), response.url)

    def test_edit_preferences(self):
        self.client.login(username='carol', password='pw12345!')
        response = self.client.post(
            reverse('edit_preferences'),
            {'target_countries': ['United Kingdom', 'Remote'], 'salary_min': '45000',
             'salary_max': '80000', 'currency': 'GBP'},
        )
        self.assertEqual(response.status_code, 302)
        prefs = UserPreferences.objects.get(user=self.user)
        self.assertEqual(prefs.target_countries, ['United Kingdom', 'Remote'])
        self.assertEqual(str(prefs.salary_min), '45000.00')
        self.assertEqual(str(prefs.salary_max), '80000.00')


@override_settings(OPENAI_API_KEY='')
class CVUploadTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username='dave', password='pw12345!')
        self.client.login(username='dave', password='pw12345!')

    def test_upload_docx_extracts_text(self):
        upload = SimpleUploadedFile(
            'jane.docx',
            build_docx_bytes(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response = self.client.post(reverse('upload_cv'), {'original_file': upload})
        self.assertEqual(response.status_code, 302)
        cv = CV.objects.get(user=self.user)
        self.assertIn('Jane Doe', cv.parsed_text)
        self.assertIn('Python, Django, SQL', cv.parsed_text)
        self.assertEqual(cv.parsed_data['raw_text'], cv.parsed_text)

    def test_upload_rejects_unsupported_type(self):
        upload = SimpleUploadedFile('notes.txt', b'hello', content_type='text/plain')
        response = self.client.post(reverse('upload_cv'), {'original_file': upload})
        self.assertEqual(response.status_code, 200)  # re-renders with error
        self.assertFalse(CV.objects.filter(user=self.user).exists())


class ApifyInputTests(TestCase):
    def test_build_actor_input_shape(self):
        from jobs.services.apify_service import _build_actor_input
        inp = _build_actor_input('United Kingdom', 30000, 50)
        # searchTerms must be a list (the actor rejects a string).
        self.assertIsInstance(inp['searchTerms'], list)
        # Country name mapped to the actor's enum code.
        self.assertEqual(inp['country'], 'uk')
        # Correct field names / minimums for this actor.
        self.assertEqual(inp['salary_min'], 30000)
        self.assertGreaterEqual(inp['max_results'], 100)
        self.assertIn('keyword', inp)
        self.assertEqual(inp['custom_location'], 'United Kingdom')
        self.assertNotIn('minSalary', inp)
        self.assertNotIn('maxItems', inp)

    def test_country_mapping_defaults_to_uk(self):
        from jobs.services.apify_service import _build_actor_input
        self.assertEqual(_build_actor_input('Narnia', None, 100)['country'], 'uk')
        self.assertEqual(_build_actor_input('United States', None, 100)['country'], 'us')

    def test_normalize_job_handles_apply_url_variant(self):
        from jobs.services.apify_service import normalize_job
        job = normalize_job({'title': 'Dev', 'company': 'Acme', 'applyUrl': 'https://x/1'})
        self.assertEqual(job['applyLink'], 'https://x/1')
        self.assertEqual(job['title'], 'Dev')


class KeywordExtractorTests(TestCase):
    def test_extract_skills_from_text(self):
        from jobs.services.keyword_extractor import extract_skills_from_text
        skills = extract_skills_from_text('Experienced in Python, Django and AWS. Also SEO.')
        self.assertIn('python', skills)
        self.assertIn('django', skills)
        self.assertIn('aws', skills)
        self.assertIn('seo', skills)
        # Word-boundary matching: "r" shouldn't match inside other words.
        self.assertNotIn('java', extract_skills_from_text('I use javascript only'))

    def test_search_keywords_prefer_job_titles(self):
        from jobs.services.keyword_extractor import extract_search_keywords
        data = {'job_titles': ['SEO Executive', 'SEO Specialist'], 'skills': ['seo']}
        self.assertEqual(extract_search_keywords(data), ['SEO Executive', 'SEO Specialist'])

    def test_search_keywords_fallback_to_raw_text(self):
        from jobs.services.keyword_extractor import extract_search_keywords
        kws = extract_search_keywords({'raw_text': 'I am a seo executive with 5 years'})
        self.assertIn('seo executive', kws)

    def test_keyword_match_score_and_missing(self):
        from jobs.services.keyword_extractor import keyword_match_score, missing_skills
        # CV covers 2 of the job's 4 required skills -> 50%.
        self.assertEqual(keyword_match_score(['python', 'sql'], ['python', 'sql', 'aws', 'go']), 50)
        # Neutral when either side has no skills (so nothing gets unfairly filtered).
        self.assertEqual(keyword_match_score([], ['python']), 50)
        self.assertEqual(missing_skills(['python'], ['python', 'aws']), ['aws'])

    def test_role_salary_range(self):
        from jobs.services.keyword_extractor import get_salary_range
        self.assertEqual(get_salary_range(['Junior Developer'])[0], 35000)  # 'developer' wins
        self.assertEqual(get_salary_range(['SEO Executive'])[0], 25000)
        self.assertEqual(get_salary_range(['Director of Ops'])[0], 60000)
        self.assertEqual(get_salary_range(['Unknown Role'], default_min=27000)[0], 27000)


class MatchingHelperTests(TestCase):
    def test_detect_sponsorship_positive(self):
        text = 'We offer visa sponsorship for skilled worker candidates.'
        self.assertEqual(detect_sponsorship(text), 'SPONSORED')

    def test_detect_sponsorship_negative(self):
        self.assertEqual(detect_sponsorship('Great team, free coffee.'), 'NOT_MENTIONED')
        self.assertEqual(detect_sponsorship(''), 'NOT_MENTIONED')

    def test_parse_salary_variants(self):
        self.assertEqual(parse_salary('£40,000 - £50,000 per annum'), 50000)
        self.assertEqual(parse_salary('45k'), 45000)
        self.assertEqual(parse_salary('Competitive'), None)
        self.assertEqual(parse_salary(''), None)

    def test_salary_threshold(self):
        # Below threshold
        meets, parsed = salary_meets_threshold('£25,000', 30000)
        self.assertFalse(meets)
        self.assertEqual(parsed, 25000)
        # Above threshold
        meets, _ = salary_meets_threshold('£60,000', 30000)
        self.assertTrue(meets)
        # Unknown salary -> included
        meets, parsed = salary_meets_threshold('Competitive', 30000)
        self.assertTrue(meets)
        self.assertIsNone(parsed)

    def test_salary_within_range(self):
        from jobs.services.matching import salary_within_range
        # Below minimum
        within, _p, reason = salary_within_range('£25,000', 30000, 60000)
        self.assertFalse(within)
        self.assertEqual(reason, 'Salary below minimum')
        # Above maximum
        within, _p, reason = salary_within_range('£90,000', 30000, 60000)
        self.assertFalse(within)
        self.assertEqual(reason, 'Salary above maximum')
        # Within range
        within, _p, reason = salary_within_range('£45,000', 30000, 60000)
        self.assertTrue(within)
        self.assertEqual(reason, '')
        # No upper limit
        within, _p, _r = salary_within_range('£200,000', 30000, None)
        self.assertTrue(within)
        # Unknown salary -> included
        within, parsed, _r = salary_within_range('Competitive', 30000, 60000)
        self.assertTrue(within)
        self.assertIsNone(parsed)

    def test_parse_match_response_clamps_and_defaults(self):
        self.assertEqual(
            _parse_match_response('{"score": 150, "reason": "great"}'),
            {'score': 100, 'reason': 'great'},
        )
        self.assertEqual(
            _parse_match_response('not json'),
            {'score': 0, 'reason': 'Unable to compute'},
        )


def _make_cv_for(user, name='Test Profile'):
    return CV.objects.create(
        user=user,
        name=name,
        original_file=SimpleUploadedFile('cv.docx', build_docx_bytes()),
        parsed_text='Python Django engineer',
    )


def fake_score(score=85, reason='Strong match', by_description=None):
    """Stand-in for tasks._score_job, which scores one job against the real CV.

    ``by_description`` maps a substring of the job description to a score, so a
    test can give each job a distinct score and prove the results are not mixed
    up across worker threads.
    """
    def _score(job_data, cv_text, use_openai):
        value, why = score, reason
        if by_description:
            for needle, mapped in by_description.items():
                if needle in (job_data.get('description') or ''):
                    value, why = mapped
                    break
        return {
            'position': job_data['position'],
            'contract': None,
            'coverage': {},
            'score': value,
            'reason': why,
        }
    return _score


TWO_RAW_JOBS = [
    {
        'title': 'Backend Engineer', 'company': 'Acme', 'location': 'London, UK',
        'datePosted': '2026-07-01', 'employmentType': 'Full-time',
        'seniorityLevel': 'Mid', 'salary': '£60,000',
        'description': 'Django REST APIs. Visa sponsorship available.',
        'applyLink': 'https://example.com/job/1',
    },
    {
        'title': 'Junior Dev', 'company': 'Beta', 'location': 'Leeds, UK',
        'datePosted': '', 'employmentType': '', 'seniorityLevel': '',
        'salary': '£20,000', 'description': 'Entry role.', 'applyLink': '',
    },
]


@override_settings(MEDIA_ROOT=_TEST_MEDIA, OPENAI_API_KEY='')
class StartSearchViewTests(TestCase):
    """The view only enqueues a task; the workflow itself is tested separately."""

    def setUp(self):
        self.user = User.objects.create_user(username='erin', password='pw12345!')
        self.client.login(username='erin', password='pw12345!')

    def test_search_requires_cv(self):
        with mock.patch('jobs.views.process_job_search.delay') as delay:
            response = self.client.post(reverse('start_search'))
        self.assertEqual(response.status_code, 302)
        self.assertEqual(SearchRun.objects.count(), 0)
        delay.assert_not_called()

    @mock.patch('jobs.views.process_job_search.delay')
    def test_search_enqueues_task_and_creates_pending_run(self, delay):
        cv = _make_cv_for(self.user)
        response = self.client.post(reverse('start_search'))
        # Redirects to dashboard, not results (async).
        self.assertRedirects(response, reverse('dashboard'))
        run = SearchRun.objects.get(user=self.user)
        self.assertEqual(run.status, SearchRun.STATUS_PENDING)
        self.assertEqual(run.cv, cv)  # search is tied to the active profile
        delay.assert_called_once_with(run.pk)

    def test_start_search_rejects_get(self):
        response = self.client.get(reverse('start_search'))
        self.assertEqual(response.status_code, 405)  # @require_POST

    def test_dashboard_renders_search_button_with_cv(self):
        _make_cv_for(self.user)
        response = self.client.get(reverse('dashboard'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Start New Search')

    def test_search_results_404_for_other_users_run(self):
        other = User.objects.create_user(username='mallory', password='pw12345!')
        run = SearchRun.objects.create(user=other)
        response = self.client.get(reverse('search_results', args=[run.pk]))
        self.assertEqual(response.status_code, 404)


class SearchStatusViewTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username='nate', password='pw12345!')
        self.client.login(username='nate', password='pw12345!')

    def test_status_returns_json(self):
        run = SearchRun.objects.create(
            user=self.user, status=SearchRun.STATUS_RUNNING, progress=40,
        )
        response = self.client.get(reverse('search_status', args=[run.pk]))
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data['status'], 'RUNNING')
        self.assertEqual(data['progress'], 40)
        self.assertEqual(data['id'], run.pk)

    def test_status_scoped_to_owner(self):
        other = User.objects.create_user(username='eve', password='pw12345!')
        run = SearchRun.objects.create(user=other)
        response = self.client.get(reverse('search_status', args=[run.pk]))
        self.assertEqual(response.status_code, 404)


@override_settings(MEDIA_ROOT=_TEST_MEDIA, OPENAI_API_KEY='')
class ProfileManagementTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username='pat', password='pw12345!')
        self.client.login(username='pat', password='pw12345!')

    def test_empty_state_when_no_profiles(self):
        response = self.client.get(reverse('dashboard'))
        self.assertContains(response, 'No CVs found')

    def test_create_profile(self):
        response = self.client.post(reverse('create_profile'), {'name': 'John Doe'})
        self.assertRedirects(response, reverse('dashboard'))
        cv = CV.objects.get(user=self.user)
        self.assertEqual(cv.name, 'John Doe')
        self.assertFalse(cv.has_file)
        # New profile becomes the active tab.
        self.assertEqual(self.client.session['active_cv_id'], cv.pk)

    def test_create_profile_rejects_blank_name(self):
        response = self.client.post(reverse('create_profile'), {'name': '   '})
        self.assertRedirects(response, reverse('dashboard'))
        self.assertEqual(CV.objects.filter(user=self.user).count(), 0)

    def test_tabs_rendered_and_active_switches_via_query(self):
        a = _make_cv_for(self.user, name='Profile A')
        b = _make_cv_for(self.user, name='Profile B')
        response = self.client.get(reverse('dashboard'))
        self.assertContains(response, 'Profile A')
        self.assertContains(response, 'Profile B')
        # Switching via ?cv_id= sets the active profile in the session.
        self.client.get(reverse('dashboard') + f'?cv_id={b.pk}')
        self.assertEqual(self.client.session['active_cv_id'], b.pk)

    def test_delete_cv_switches_active_and_removes_record(self):
        a = _make_cv_for(self.user, name='Profile A')
        b = _make_cv_for(self.user, name='Profile B')
        # Make B active, then delete it -> falls back to remaining profile.
        self.client.get(reverse('dashboard') + f'?cv_id={b.pk}')
        response = self.client.post(reverse('delete_cv', args=[b.pk]))
        self.assertRedirects(response, reverse('dashboard'))
        self.assertFalse(CV.objects.filter(pk=b.pk).exists())
        self.assertEqual(self.client.session['active_cv_id'], a.pk)

    def test_delete_other_users_cv_forbidden(self):
        other = User.objects.create_user(username='intruder', password='pw12345!')
        cv = _make_cv_for(other, name='Secret')
        response = self.client.post(reverse('delete_cv', args=[cv.pk]))
        self.assertEqual(response.status_code, 404)
        self.assertTrue(CV.objects.filter(pk=cv.pk).exists())

    def test_upload_targets_active_profile(self):
        cv = _make_cv_for(self.user, name='Profile A')
        cv.original_file.delete(save=False)  # empty the profile
        cv.original_file = ''
        cv.parsed_text = ''
        cv.save()
        self.client.get(reverse('dashboard') + f'?cv_id={cv.pk}')  # make active

        upload = SimpleUploadedFile('resume.docx', build_docx_bytes())
        response = self.client.post(reverse('upload_cv'), {'original_file': upload})
        self.assertRedirects(response, reverse('dashboard'))
        # Same profile updated, not a new CV created.
        self.assertEqual(CV.objects.filter(user=self.user).count(), 1)
        cv.refresh_from_db()
        self.assertTrue(cv.has_file)
        self.assertIn('Jane Doe', cv.parsed_text)


class PreferencesCurrencyTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username='cara', password='pw12345!')
        self.client.login(username='cara', password='pw12345!')

    def test_currency_saved_and_default_gbp(self):
        response = self.client.post(
            reverse('edit_preferences'),
            {'target_countries': ['United Kingdom'], 'salary_min': '50000', 'currency': 'USD'},
        )
        self.assertRedirects(response, reverse('dashboard'))
        prefs = UserPreferences.objects.get(user=self.user)
        self.assertEqual(prefs.currency, 'USD')
        self.assertEqual(prefs.currency_symbol, '$')

    def test_salary_max_must_exceed_min(self):
        response = self.client.post(
            reverse('edit_preferences'),
            {'target_countries': [], 'salary_min': '60000', 'salary_max': '40000', 'currency': 'GBP'},
        )
        self.assertEqual(response.status_code, 200)  # re-rendered with error
        self.assertFalse(UserPreferences.objects.filter(user=self.user, salary_min=60000).exists())


class ClearHistoryAndStatusTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username='hank', password='pw12345!')
        self.client.login(username='hank', password='pw12345!')

    def test_clear_search_history(self):
        SearchRun.objects.create(user=self.user, status=SearchRun.STATUS_COMPLETED)
        SearchRun.objects.create(user=self.user, status=SearchRun.STATUS_FAILED)
        other = User.objects.create_user(username='someoneelse', password='pw12345!')
        keep = SearchRun.objects.create(user=other)

        response = self.client.post(reverse('clear_search_history'))
        self.assertRedirects(response, reverse('dashboard'))
        self.assertEqual(SearchRun.objects.filter(user=self.user).count(), 0)
        self.assertTrue(SearchRun.objects.filter(pk=keep.pk).exists())  # other user's kept

    def test_clear_history_rejects_get(self):
        response = self.client.get(reverse('clear_search_history'))
        self.assertEqual(response.status_code, 405)

    def test_status_returns_total_and_processed(self):
        run = SearchRun.objects.create(
            user=self.user, status=SearchRun.STATUS_RUNNING, progress=50, total_jobs=10,
        )
        Job.objects.create(search_run=run, title='A', company='X', location='L',
                           application_link='https://x/1')
        data = self.client.get(reverse('search_status', args=[run.pk])).json()
        self.assertEqual(data['total'], 10)
        self.assertEqual(data['processed'], 1)
        self.assertEqual(data['progress'], 50)


@override_settings(MEDIA_ROOT=_TEST_MEDIA, OPENAI_API_KEY='')
class RunJobSearchTests(TestCase):
    """Exercises the async workflow function directly (no Celery/broker needed)."""

    def setUp(self):
        self.user = User.objects.create_user(username='erin', password='pw12345!')
        _make_cv_for(self.user)
        self.run = SearchRun.objects.create(
            user=self.user, countries=['United Kingdom'], min_salary=30000,
            status=SearchRun.STATUS_PENDING,
        )

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks._score_job')
    @mock.patch('jobs.tasks.search_jobs')
    def test_workflow_creates_jobs_and_completes(self, mock_search, mock_score, mock_sheet):
        from jobs.tasks import run_job_search
        mock_search.return_value = TWO_RAW_JOBS
        mock_score.side_effect = fake_score(85, 'Strong match')

        result = run_job_search(self.run.pk)
        self.run.refresh_from_db()

        self.assertEqual(result['status'], 'COMPLETED')
        self.assertEqual(self.run.status, SearchRun.STATUS_COMPLETED)
        self.assertEqual(self.run.progress, 100)

        jobs = {j.title: j for j in self.run.jobs.all()}
        self.assertEqual(jobs['Backend Engineer'].match_score, 85)
        self.assertEqual(jobs['Backend Engineer'].sponsorship_flag, 'SPONSORED')
        self.assertTrue(jobs['Backend Engineer'].tailored_pdf)
        self.assertTrue(jobs['Backend Engineer'].processed)

        # The £20k job is under the run's minimum, and it is still here, still
        # scored, and still tailored. The minimum is a hint to the Apify actor, not
        # a filter we re-apply to bin what came back — an advert's figure is a
        # range, and the candidate can decide for themselves.
        self.assertEqual(jobs['Junior Dev'].match_score, 85)
        self.assertNotIn('below minimum', jobs['Junior Dev'].match_reason)
        self.assertFalse(jobs['Junior Dev'].above_salary_preference)
        self.assertTrue(jobs['Junior Dev'].tailored_pdf)

        self.assertEqual(mock_score.call_count, 2)  # every job scored, not just one
        # Every job is logged to the candidate's Google Sheets tab.
        self.assertEqual(mock_sheet.return_value.log_job.call_count, 2)

    @mock.patch('jobs.tasks.search_jobs')
    def test_workflow_marks_failed_on_apify_error(self, mock_search):
        from jobs.services.apify_service import ApifySearchError
        from jobs.tasks import run_job_search
        mock_search.side_effect = ApifySearchError('boom')

        result = run_job_search(self.run.pk)
        self.run.refresh_from_db()
        self.assertEqual(result['status'], 'FAILED')
        self.assertEqual(self.run.status, SearchRun.STATUS_FAILED)
        self.assertIn('boom', self.run.error_message)
        self.assertEqual(self.run.jobs.count(), 0)

    @override_settings(OPENAI_MAX_SCORED_JOBS=1)
    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks._score_job')
    @mock.patch('jobs.tasks.search_jobs')
    def test_cap_limits_precision_not_whether_a_job_is_scored(
        self, mock_search, mock_score, mock_sheet
    ):
        """The cost cap must never leave a job with a placeholder score.

        It used to: everything past OPENAI_MAX_SCORED_JOBS was stored with a
        hardcoded 0 and "Not scored", which made the scores meaningless and buried
        good jobs. The cap now decides only HOW a job is scored — model-extracted
        contract for the best-ranked, deterministic contract for the rest — not
        WHETHER it is scored.
        """
        from jobs.tasks import run_job_search
        mock_search.return_value = [
            {'title': 'A', 'company': 'X', 'location': 'London', 'salary': '£60,000',
             'description': 'Role A', 'applyLink': '', 'datePosted': '',
             'employmentType': '', 'seniorityLevel': ''},
            {'title': 'B', 'company': 'Y', 'location': 'London', 'salary': '£60,000',
             'description': 'Role B', 'applyLink': '', 'datePosted': '',
             'employmentType': '', 'seniorityLevel': ''},
        ]
        mock_score.side_effect = fake_score(50, 'ok')

        run_job_search(self.run.pk)

        # Both jobs scored...
        self.assertEqual(mock_score.call_count, 2)
        jobs = list(self.run.jobs.all())
        self.assertEqual(len(jobs), 2)
        for job in jobs:
            self.assertEqual(job.match_score, 50)
            self.assertNotIn('Not scored', job.match_reason)

        # ...but only one of them with a model-extracted contract.
        precise = [call.args[2] for call in mock_score.call_args_list]
        self.assertEqual(sorted(precise), [False, True])


@override_settings(MEDIA_ROOT=_TEST_MEDIA, CELERY_TASK_ALWAYS_EAGER=True, OPENAI_API_KEY='')
class CeleryTaskTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username='oscar', password='pw12345!')
        _make_cv_for(self.user)
        self.run = SearchRun.objects.create(
            user=self.user, countries=['United Kingdom'], status=SearchRun.STATUS_PENDING,
        )

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks._score_job')
    @mock.patch('jobs.tasks.search_jobs')
    def test_task_runs_workflow(self, mock_search, mock_score, mock_sheet):
        from jobs.tasks import process_job_search
        mock_search.return_value = [TWO_RAW_JOBS[0]]
        mock_score.side_effect = fake_score(80, 'ok')

        # Call the task function directly (equivalent to eager execution).
        result = process_job_search.run(self.run.pk)
        self.assertEqual(result['status'], 'COMPLETED')
        self.run.refresh_from_db()
        self.assertEqual(self.run.status, SearchRun.STATUS_COMPLETED)

    def test_task_handles_missing_run(self):
        from jobs.tasks import process_job_search
        result = process_job_search.run(999999)
        self.assertEqual(result['status'], 'FAILED')


class TailoringTests(TestCase):
    def test_returns_original_when_openai_not_configured(self):
        # With no OPENAI_API_KEY configured, tailoring falls back to the original.
        with override_settings(OPENAI_API_KEY=''):
            result = tailor_cv_for_job('My CV text', 'Job desc', 'Engineer', 'Acme')
        self.assertEqual(result, 'My CV text')

    def test_empty_cv_returns_empty(self):
        self.assertEqual(tailor_cv_for_job('', 'Job desc', 'Engineer', 'Acme'), '')


class PdfGeneratorTests(TestCase):
    def test_build_pdf_filename_sanitizes(self):
        name = build_pdf_filename('Jane Doe', 'Senior Dev/Eng', 'Acme, Inc.')
        self.assertEqual(name, 'Jane_Doe_Senior_Dev_Eng_Acme__Inc_.pdf')

    def test_generates_valid_pdf_file(self):
        path = os.path.join(_TEST_MEDIA, 'out.pdf')
        cv_text = (
            'Summary\nExperienced engineer.\n'
            'Experience\nBuilt APIs at Acme.\n'
            'Skills\nPython, Django.'
        )
        generate_tailored_pdf(cv_text, 'Jane Doe', 'Engineer', 'Acme', path)
        self.assertTrue(os.path.exists(path))
        self.assertGreater(os.path.getsize(path), 0)
        with open(path, 'rb') as fh:
            self.assertEqual(fh.read(5), b'%PDF-')  # PDF magic bytes

    def test_generates_pdf_without_headings(self):
        path = os.path.join(_TEST_MEDIA, 'plain.pdf')
        generate_tailored_pdf('Just some plain text.\nNo headings here.',
                              'John', 'Dev', 'Beta', path)
        with open(path, 'rb') as fh:
            self.assertEqual(fh.read(5), b'%PDF-')


UK_CV_TEXT = """AMANDA TURNER
07123 456789 | amanda@example.com | London, UK

PROFESSIONAL PROFILE
Results-driven SEO Executive with 4+ years of experience in organic search and
content strategy, seeking a role in a dynamic agency.

KEY SKILLS
- Technical SEO
- Google Analytics
- Keyword Research
- SEMrush & Ahrefs

PROFESSIONAL EXPERIENCE
SEO Executive | XYZ Agency | London, UK
Jan 2022 - Present
- Led SEO strategy for 25+ clients, increasing organic traffic by 150%.
- Managed a team of 3 junior SEOs.

EDUCATION
BSc Digital Marketing | University of Manchester | Manchester, UK
2017 - 2020
- 2:1 (Upper Second Class Honours)

CERTIFICATIONS
- Google Analytics Individual Qualification
"""


class UkCvFormatTests(TestCase):
    def test_parses_all_sections(self):
        from jobs.services.pdf_generator import parse_cv_sections
        s = parse_cv_sections(UK_CV_TEXT)
        self.assertEqual(s['name'], 'AMANDA TURNER')
        self.assertIn('amanda@example.com', s['contact'])
        self.assertEqual(len(s['skills']), 4)
        self.assertIn('SEO Executive | XYZ Agency | London, UK', s['experience'])
        self.assertTrue(s['education'])
        self.assertTrue(s['certifications'])

    def test_body_text_containing_section_words_is_not_a_heading(self):
        """Regression: 'years of experience' must not be read as the EXPERIENCE heading."""
        from jobs.services.pdf_generator import parse_cv_sections
        s = parse_cv_sections(UK_CV_TEXT)
        profile = ' '.join(s['profile'])
        # The whole profile survives, including the sentence after 'experience'.
        self.assertIn('seeking a role in a dynamic agency', profile)
        # And the experience section holds only real roles, not profile prose.
        self.assertNotIn('Results-driven', ' '.join(s['experience']))

    def test_heading_aliases_and_markdown(self):
        from jobs.services.pdf_generator import parse_cv_sections
        s = parse_cv_sections(
            'Bob\n\n## Personal Statement\nHi.\n\n**Work History**\nAcme | Dev\n'
        )
        self.assertIn('Hi.', s['profile'])
        self.assertIn('Acme | Dev', s['experience'])

    def test_renders_uk_pdf_with_sections(self):
        from PyPDF2 import PdfReader
        path = os.path.join(_TEST_MEDIA, 'uk.pdf')
        generate_tailored_pdf(UK_CV_TEXT, 'Amanda Turner', 'SEO Executive', 'XYZ', path)
        text = PdfReader(path).pages[0].extract_text()
        for needle in ['AMANDA TURNER', 'PROFESSIONAL PROFILE', 'KEY SKILLS',
                       'PROFESSIONAL EXPERIENCE', 'EDUCATION', 'CERTIFICATIONS',
                       'Technical SEO', 'Upper Second Class']:
            self.assertIn(needle, text)


ATS_JOB_DESCRIPTION = """Senior Backend Engineer — London, UK

We are looking for a Senior Backend Engineer with 5+ years of experience building
Python and Django services. You will design REST APIs, work with PostgreSQL and
SQL, and deploy to AWS using Docker and Kubernetes.

Requirements:
- 5+ years experience in commercial software development
- Bachelor's degree in Computer Science or similar
- Strong Python and Django skills
- Experience with AWS, Docker, Kubernetes and PostgreSQL
- Excellent communication and leadership skills
- Right to work in the UK (we do not offer visa sponsorship)
"""

# A strong CV: right skills, dates, quantified bullets, standard headings.
ATS_GOOD_CV = """JOHN SMITH
07123 456789 | john@example.com | London, UK

PROFESSIONAL PROFILE
Senior Backend Engineer with 7 years of experience building Python and Django
services. Skilled in AWS, Docker, Kubernetes and PostgreSQL, with strong
communication and leadership.

KEY SKILLS
- Python
- Django
- PostgreSQL
- AWS
- Docker
- Kubernetes

PROFESSIONAL EXPERIENCE
Senior Backend Engineer | Acme Ltd | London, UK
Jan 2021 - Present
- Led a team of 5 engineers building Django REST APIs, cutting latency by 40%.
- Migrated 30 services to Docker and Kubernetes on AWS, saving £250,000 a year.
- Optimised PostgreSQL queries with SQL tuning, improving throughput by 3x.

Backend Engineer | Beta Corp | Manchester, UK
Feb 2019 - Dec 2020
- Built Python microservices handling 2 million requests per day.

EDUCATION
BSc Computer Science | University of Manchester | Manchester, UK
2015 - 2018
- First Class Honours

CERTIFICATIONS
- AWS Certified Solutions Architect
"""

# A weak CV: no standard headings, no dates, no numbers, almost no keywords.
ATS_BAD_CV = """Bob Jones

My Journey
I have done some work with computers over the years and enjoy solving problems.

Stuff I Can Do
Microsoft Word, answering the phone, being punctual.
"""


class ATSTextHelperTests(TestCase):
    def test_stemming_unifies_word_forms(self):
        from jobs.services.ats_checker import _stem
        stems = {_stem(w) for w in ['manage', 'managing', 'managed', 'management']}
        self.assertEqual(len(stems), 1, f'expected one stem, got {stems}')
        # Short acronyms must survive untouched.
        self.assertEqual(_stem('sql'), 'sql')
        self.assertEqual(_stem('aws'), 'aws')

    def test_synonyms_are_symmetric(self):
        from jobs.services.ats_checker import _synonyms_for
        self.assertIn('k8s', _synonyms_for('kubernetes'))
        self.assertIn('kubernetes', _synonyms_for('k8s'))

    def test_total_experience_merges_overlapping_roles(self):
        from jobs.services.ats_checker import total_experience_years
        # Two concurrent roles over the same 2 years must not count as 4.
        text = 'Jan 2020 - Jan 2022\nJun 2020 - Jan 2022'
        self.assertEqual(total_experience_years(text), 2.0)

    def test_total_experience_sums_sequential_roles(self):
        from jobs.services.ats_checker import total_experience_years
        self.assertEqual(
            total_experience_years('Jan 2018 - Jan 2020\nJan 2021 - Jan 2022'), 3.0
        )

    def test_no_dates_yields_zero(self):
        from jobs.services.ats_checker import total_experience_years
        self.assertEqual(total_experience_years('I have loads of experience'), 0.0)


class ATSJobContextTests(TestCase):
    def test_carries_title_and_location(self):
        from jobs.services.ats_checker import extract_job_context
        context = extract_job_context(ATS_JOB_DESCRIPTION, 'Senior Backend Engineer',
                                      'London, UK')
        self.assertEqual(context['title'], 'Senior Backend Engineer')
        self.assertEqual(context['location'], 'London, UK')

    def test_mines_no_hard_requirements(self):
        """The requirement mining is gone, and with it every knock-out it fed.

        An advert demanding "5 years" and a "PMP" used to produce filters that
        binned the candidate outright. Nothing is mined to filter on any more.
        """
        from jobs.services.ats_checker import extract_job_context
        context = extract_job_context(
            ATS_JOB_DESCRIPTION + '\n10 years experience and a PMP are required.'
        )
        self.assertEqual(set(context), {'title', 'location'})

    def test_keywords_rank_hard_skills_above_soft(self):
        from jobs.services.ats_checker import extract_jd_keywords
        keywords = {k['term']: k for k in extract_jd_keywords(ATS_JOB_DESCRIPTION)}
        self.assertIn('python', keywords)
        self.assertEqual(keywords['python']['type'], 'hard')
        self.assertGreater(
            keywords['python']['weight'], keywords['communication']['weight']
        )


class NoKnockoutTests(TestCase):
    """The checker scores. It never rejects — whatever the CV or the advert says.

    Each case below used to be a hard knock-out that binned the job outright.
    Every one of them must now come back with a real score and no rejection, and
    the report must carry no rejection machinery at all for a caller to trip over.
    """

    def _report(self, cv_text, job=ATS_JOB_DESCRIPTION, location='London, UK'):
        from jobs.services.ats_checker import check_cv_against_job
        return check_cv_against_job(
            cv_text, job, 'Senior Backend Engineer', location,
        )

    def test_report_carries_no_rejection_verdict(self):
        report = self._report(ATS_GOOD_CV)
        for key in ('rejected', 'rejection_reasons', 'knockout_reasons'):
            self.assertNotIn(key, report)
        self.assertNotIn('phase2_knockout', report['phases'])

    def test_too_little_experience_still_scores(self):
        cv = ATS_GOOD_CV.replace('Jan 2021 - Present', 'Jan 2024 - Present') \
                        .replace('Feb 2019 - Dec 2020', 'Feb 2023 - Dec 2023')
        report = self._report(cv)
        self.assertIsInstance(report['overall_score'], int)
        self.assertGreater(report['overall_score'], 0)

    def test_missing_degree_still_scores(self):
        cv = ATS_GOOD_CV.replace('BSc Computer Science', 'Evening course in coding')
        self.assertGreater(self._report(cv)['overall_score'], 0)

    def test_missing_certification_still_scores(self):
        job = ATS_JOB_DESCRIPTION + '\nYou must hold a valid PMP certification.'
        self.assertGreater(self._report(ATS_GOOD_CV, job=job)['overall_score'], 0)

    def test_conflicting_location_still_scores(self):
        """The knock-out that hurt most: a candidate in the wrong city, binned."""
        cv = ATS_GOOD_CV.replace('London, UK', 'Edinburgh, Scotland')
        self.assertGreater(self._report(cv)['overall_score'], 0)

    def test_cv_requiring_sponsorship_still_scores(self):
        cv = ATS_GOOD_CV + '\nI require sponsorship to work in the UK.'
        self.assertGreater(self._report(cv)['overall_score'], 0)

    def test_unparsable_cv_still_scores(self):
        """An 'unparsable CV' is a formatting note, not a verdict on 200 jobs."""
        report = self._report('John Smith\nPython Django engineer')
        self.assertIsInstance(report['overall_score'], int)
        self.assertNotIn('rejected', report)


class ATSScoringTests(TestCase):
    def _report(self, cv_text, **kwargs):
        from jobs.services.ats_checker import check_cv_against_job
        return check_cv_against_job(
            cv_text, ATS_JOB_DESCRIPTION,
            kwargs.get('title', 'Senior Backend Engineer'),
            kwargs.get('location', 'London, UK'),
        )

    def test_strong_cv_scores_well_and_passes(self):
        report = self._report(ATS_GOOD_CV)
        self.assertGreaterEqual(report['overall_score'], 75)
        self.assertTrue(report['pass'])

    def test_weak_cv_scores_low(self):
        report = self._report(ATS_BAD_CV)
        self.assertLess(report['overall_score'], 50)
        self.assertFalse(report['pass'])
        self.assertTrue(report['recommendations'])

    def test_missing_sections_are_advice_not_a_rejection(self):
        """Creative headings ('My Journey') are hard to parse — and nothing more."""
        report = self._report(ATS_BAD_CV)
        phase1 = report['phases']['phase1_parsing']
        self.assertIn('Experience', phase1['missing_headers'])
        self.assertIn('My Journey', phase1['creative_headers'])
        # Reported as an advisory formatting issue; the job is still scored, and
        # the report has no rejection verdict to act on.
        self.assertTrue(report['formatting_issues'])
        self.assertNotIn('rejected', report)

    def test_keyword_phase_finds_and_misses_the_right_skills(self):
        cv = ATS_GOOD_CV.replace('- Kubernetes\n', '').replace(
            'Docker and Kubernetes on AWS', 'AWS'
        ).replace('Skilled in AWS, Docker, Kubernetes and PostgreSQL',
                  'Skilled in AWS and PostgreSQL')
        phase3 = self._report(cv)['phases']['phase3_keyword']
        self.assertIn('python', phase3['hard_skills_found'])
        self.assertIn('kubernetes', phase3['hard_skills_missing'])

    def test_synonyms_count_as_a_match(self):
        """'k8s' on the CV must satisfy the JD's 'kubernetes'."""
        cv = ATS_GOOD_CV.replace('Kubernetes', 'k8s')
        phase3 = self._report(cv)['phases']['phase3_keyword']
        self.assertIn('kubernetes', phase3['hard_skills_found'])

    def test_keyword_stuffing_is_penalised(self):
        stuffed = ATS_GOOD_CV + '\n' + ('Python Django Kubernetes. ' * 30)
        report = self._report(stuffed)
        self.assertTrue(report['phases']['phase3_keyword']['keyword_stuffing'])

    def test_quantification_rewards_numbers(self):
        phase4 = self._report(ATS_GOOD_CV)['phases']['phase4_context']
        # Every bullet in the good CV carries a figure.
        self.assertGreaterEqual(phase4['quantification_score'], 75)

    def test_skills_listed_but_never_evidenced_score_low_on_proximity(self):
        cv = """JANE DOE
London, UK

PROFESSIONAL PROFILE
An engineer.

KEY SKILLS
- Python
- Django
- Kubernetes
- Docker
- AWS
- PostgreSQL

PROFESSIONAL EXPERIENCE
Engineer | Acme | London, UK
Jan 2016 - Present
- Did various tasks.

EDUCATION
BSc Computer Science | Leeds | 2012 - 2015
"""
        phase4 = self._report(cv)['phases']['phase4_context']
        self.assertLess(phase4['proximity_score'], 40)

    def test_chronology_detects_gaps_and_ordering(self):
        from jobs.services.ats_checker import ATSChecker
        cv = """JOHN SMITH

PROFESSIONAL EXPERIENCE
Junior Dev | Old Co | London
Jan 2015 - Jan 2017

Senior Dev | New Co | London
Jan 2019 - Jan 2021

EDUCATION
BSc Computing | Leeds | 2012 - 2015
"""
        checker = ATSChecker(cv, ATS_JOB_DESCRIPTION)
        phase5 = checker.check_chronology()
        # Oldest role listed first -> not reverse-chronological.
        self.assertFalse(phase5['reverse_chronological'])
        # A 24-month hole between the two roles.
        self.assertEqual(len(phase5['gaps']), 1)
        self.assertEqual(phase5['gaps'][0]['months'], 24)

    def test_education_parsing(self):
        phase6 = self._report(ATS_GOOD_CV)['phases']['phase6_education']
        self.assertEqual(phase6['degree_hierarchy'], "Bachelor's")
        self.assertEqual(phase6['graduation_year'], 2018)
        self.assertEqual(phase6['gpa'], 4.0)  # First Class Honours

    def test_report_has_the_documented_shape(self):
        report = self._report(ATS_GOOD_CV)
        for key in ('overall_score', 'ats_score', 'pass', 'phases',
                    'sectional_scores', 'recommendations', 'text_hash'):
            self.assertIn(key, report)
        # No phase 2: the knock-out phase is gone, and the numbering keeps its gap.
        for phase in ('phase1_parsing', 'phase3_keyword',
                      'phase4_context', 'phase5_experience', 'phase6_education'):
            self.assertIn(phase, report['phases'])
        self.assertEqual(
            set(report['sectional_scores']), {'skills', 'experience', 'education'}
        )
        self.assertTrue(0 <= report['overall_score'] <= 100)

    def test_deduplication_hash_ignores_whitespace(self):
        from jobs.services.ats_checker import ATSChecker
        a = ATSChecker('Hello   World', '').text_hash()
        b = ATSChecker('hello world', '').text_hash()
        self.assertEqual(a, b)


class ATSProximityAndFrequencyTests(TestCase):
    """Phase 3 frequency and Phase 4 proximity, per the ATS methodology."""

    def test_words_of_a_requirement_must_sit_together_in_one_bullet(self):
        from jobs.services.ats_checker import _within_window
        bullet = _tokens_of('- Owned financial forecasting for 12 regions.')
        self.assertTrue(_within_window(bullet, ['financial', 'forecasting']))

    def test_scattered_words_do_not_count_as_proximity(self):
        from jobs.services.ats_checker import _within_window
        scattered = _tokens_of(
            'Financial services background with a long record of delivery, plus '
            'planning, budgeting, reporting, analysis and long-range forecasting'
        )
        # Both words present, but far apart -> not used together.
        self.assertFalse(_within_window(scattered, ['financial', 'forecasting']))

    def test_proximity_scores_a_real_bullet_above_a_scattered_cv(self):
        from jobs.services.ats_checker import ATSChecker
        job = ('We need financial forecasting and financial forecasting skills. '
               'Financial forecasting is core to this role.')
        together = """JANE DOE

PROFESSIONAL EXPERIENCE
Analyst | Acme | London
Jan 2018 - Present
- Led financial forecasting across 12 regions, improving accuracy by 30%.

EDUCATION
BSc Finance | Leeds | 2014 - 2017
"""
        apart = """JANE DOE

PROFESSIONAL EXPERIENCE
Analyst | Acme | London
Jan 2018 - Present
- Worked in financial services for many clients across the region.
- Handled forecasting duties as required by the wider team each quarter.

EDUCATION
BSc Finance | Leeds | 2014 - 2017
"""
        near = ATSChecker(together, job).run_phase4()['proximity_score']
        far = ATSChecker(apart, job).run_phase4()['proximity_score']
        self.assertGreater(near, far)

    def test_expected_frequency_follows_the_jd(self):
        from jobs.services.ats_checker import _expected_frequency
        # "If the JD mentions SQL 5 times, your CV needs 2-3 occurrences."
        self.assertIn(_expected_frequency(5), (2, 3))
        self.assertEqual(_expected_frequency(1), 1)
        self.assertEqual(_expected_frequency(20), 3)  # capped: more is stuffing

    def test_underused_keyword_is_flagged(self):
        from jobs.services.ats_checker import check_cv_against_job
        job = 'SQL SQL SQL SQL SQL. We need SQL and Python.'
        cv = """JOHN SMITH

PROFESSIONAL PROFILE
Engineer.

KEY SKILLS
- SQL
- Python

PROFESSIONAL EXPERIENCE
Engineer | Acme | London
Jan 2018 - Present
- Built things with Python.

EDUCATION
BSc Computing | Leeds | 2014 - 2017
"""
        phase3 = check_cv_against_job(cv, job)['phases']['phase3_keyword']
        underused = {u['term']: u for u in phase3['underused_keywords']}
        self.assertIn('sql', underused)
        self.assertEqual(underused['sql']['cv_count'], 1)
        self.assertGreaterEqual(underused['sql']['expected'], 2)


def _tokens_of(text):
    from jobs.services.ats_checker import _tokens
    return _tokens(text)


class ATSGradeReadingTests(TestCase):
    """Grades are read to SCORE the education section, never to filter on."""

    def _gpa(self, cv):
        from jobs.services.ats_checker import ATSChecker
        return ATSChecker(cv, ATS_JOB_DESCRIPTION)._cv_gpa()

    def test_gpa_is_read_from_the_cv(self):
        self.assertEqual(self._gpa(ATS_GOOD_CV.replace('First Class Honours', 'GPA 3.9')),
                         3.9)

    def test_uk_classification_maps_onto_the_gpa_scale(self):
        self.assertEqual(self._gpa(ATS_GOOD_CV), 4.0)  # First Class Honours
        self.assertEqual(
            self._gpa(ATS_GOOD_CV.replace('First Class Honours',
                                          'Lower Second Class (2:2)')),
            3.0,
        )

    def test_a_low_grade_never_bins_the_job(self):
        from jobs.services.ats_checker import check_cv_against_job
        cv = ATS_GOOD_CV.replace('First Class Honours', 'GPA 2.8')
        report = check_cv_against_job(cv, 'Minimum GPA 3.5 required. Bachelor degree.')
        self.assertNotIn('rejected', report)
        self.assertGreater(report['overall_score'], 0)


class ATSGarbledFontTests(TestCase):
    def test_garbled_text_is_flagged_as_a_font_issue(self):
        from jobs.services.ats_checker import ATSChecker
        garbled = ATS_GOOD_CV.replace('Python', 'P�th�n').replace(
            'Django', 'Dj�ng�'
        )
        checker = ATSChecker(garbled, ATS_JOB_DESCRIPTION)
        # No file to inspect, so drive the text-level check directly.
        self.assertGreater(checker._garbled_ratio(), 0.001)

    def test_clean_text_is_not_flagged(self):
        from jobs.services.ats_checker import ATSChecker
        self.assertEqual(ATSChecker(ATS_GOOD_CV, '')._garbled_ratio(), 0.0)


class ATSKeywordExtractionQualityTests(TestCase):
    """The keyword list must contain things a CV can actually match."""

    REAL_AD = """Senior Backend Engineer - London

About Us
We are a fast-growing fintech. We are proud to be an equal opportunity employer.

What We're Looking For
- 5+ years of commercial software engineering experience
- Strong Python skills and deep experience with Django
- Experience with PostgreSQL and optimising SQL queries
- Familiarity with Docker and Kubernetes

Benefits
- Competitive salary and equity
- 25 days holiday plus bank holidays
- Private healthcare, and we offer a learning budget
"""

    def test_boilerplate_sections_are_not_mined_for_keywords(self):
        from jobs.services.ats_checker import extract_jd_keywords
        terms = {k['term'] for k in extract_jd_keywords(self.REAL_AD)}
        # Benefits and culture blurb must contribute no keywords: no CV can or
        # should match "holiday", "equity" or "healthcare".
        for noise in ('holiday', 'equity', 'healthcare', 'opportunity', 'salary'):
            self.assertNotIn(noise, terms)

    def test_real_skills_are_still_extracted(self):
        from jobs.services.ats_checker import extract_jd_keywords
        terms = {k['term'] for k in extract_jd_keywords(self.REAL_AD)}
        for skill in ('python', 'django', 'postgresql', 'sql', 'docker', 'kubernetes'):
            self.assertIn(skill, terms)

    def test_verbs_and_adjectives_are_not_demanded_as_keywords(self):
        from jobs.services.ats_checker import extract_jd_keywords
        jd = ('You will design and build scalable services. You will build and '
              'design and deploy and maintain systems using Python. Deploy often.')
        terms = {k['term'] for k in extract_jd_keywords(jd)}
        # "using", "deploy", "build", "design" are prose, not requirements.
        for verb in ('using', 'build', 'design'):
            self.assertNotIn(verb, terms)
        self.assertIn('python', terms)

    def test_placement_bonus_cannot_mask_a_missing_keyword(self):
        """A CV missing a hard skill must never score 100 on keywords."""
        from jobs.services.ats_checker import check_cv_against_job
        cv = ATS_GOOD_CV.replace('Kubernetes', '').replace('- Kubernetes\n', '')
        phase3 = check_cv_against_job(cv, ATS_JOB_DESCRIPTION)['phases']['phase3_keyword']
        self.assertIn('kubernetes', phase3['hard_skills_missing'])
        self.assertLess(phase3['score'], 100)


ORIGINAL_CV_NO_NUMBERS = """John Smith
07123 456789 | john@example.com | London, UK

About Me
A software developer who enjoys building things.

Employment
Software Developer, Acme Ltd, London
January 2021 - Present
- Worked on the back end and made database queries run faster.
- Helped move applications onto the cloud using containers.

Education
BSc Computer Science, University of Manchester, 2015 - 2018

Technical
Python, Django, Postgres, Docker, Amazon Web Services
"""


class ATSFabricationGuardrailTests(TestCase):
    """The tailored CV must never claim what the original cannot back up.

    Prompt instructions do not reliably prevent this — a live run had the model
    invent 'CI/CD', 'improving performance by 30%' and 'over 10,000 users' to
    chase the keyword and quantification scores. So the output is verified
    against the source rather than trusted.
    """

    JD = ('Backend Engineer. We need Python, Django, PostgreSQL, Docker, '
          'Kubernetes and CI/CD experience. 3+ years.')

    def test_invented_skill_is_detected(self):
        from jobs.services.ats_checker import unsupported_claims
        tailored = ORIGINAL_CV_NO_NUMBERS + '\nKEY SKILLS\n- CI/CD\n- Kubernetes\n'
        claims = unsupported_claims(ORIGINAL_CV_NO_NUMBERS, tailored, self.JD)
        self.assertIn('ci/cd', claims)
        self.assertIn('kubernetes', claims)

    def test_implied_skills_are_not_false_positives(self):
        """'Postgres' on the original CV genuinely evidences a 'SQL' claim."""
        from jobs.services.ats_checker import unsupported_claims
        tailored = ORIGINAL_CV_NO_NUMBERS + '\n- Optimised SQL queries.\n'
        claims = unsupported_claims(ORIGINAL_CV_NO_NUMBERS, tailored, self.JD)
        self.assertNotIn('sql', claims)

    def test_genuine_skills_are_not_flagged(self):
        from jobs.services.ats_checker import unsupported_claims
        tailored = ORIGINAL_CV_NO_NUMBERS + '\n- Built services in Python and Django.\n'
        claims = unsupported_claims(ORIGINAL_CV_NO_NUMBERS, tailored, self.JD)
        self.assertNotIn('python', claims)
        self.assertNotIn('django', claims)

    def test_invented_metrics_are_detected(self):
        from jobs.services.ats_checker import fabricated_metrics
        tailored = (ORIGINAL_CV_NO_NUMBERS
                    + '\n- Improved performance by 30% for over 10,000 users.\n')
        invented = fabricated_metrics(ORIGINAL_CV_NO_NUMBERS, tailored)
        self.assertTrue(any('30' in m for m in invented))
        self.assertTrue(any('10,000' in m or '10000' in m for m in invented))

    def test_metrics_present_in_the_original_are_kept(self):
        from jobs.services.ats_checker import fabricated_metrics
        original = ORIGINAL_CV_NO_NUMBERS + '\n- Cut latency by 40%.\n'
        tailored = original.replace('Cut latency by 40%', 'Reduced latency by 40%')
        self.assertEqual(fabricated_metrics(original, tailored), [])

    def test_identical_text_never_flags_anything(self):
        """Regression: magnitudes ("2 million") were parsed differently on each
        side of the comparison, so a CV was flagged for inventing its own numbers.
        """
        from jobs.services.ats_checker import fabricated_metrics, unsupported_claims
        self.assertEqual(fabricated_metrics(ATS_GOOD_CV, ATS_GOOD_CV), [])
        self.assertEqual(
            unsupported_claims(ATS_GOOD_CV, ATS_GOOD_CV, ATS_JOB_DESCRIPTION), []
        )

    def test_magnitudes_and_currency_are_compared_by_value(self):
        from jobs.services.ats_checker import fabricated_metrics
        original = 'Handled 2 million requests and saved £250,000.'
        # Same figures, reworded -> nothing invented.
        self.assertEqual(
            fabricated_metrics(original, 'Saved £250,000 while serving 2 million users.'),
            [],
        )
        # A different figure -> invented.
        self.assertTrue(fabricated_metrics(original, 'Handled 9 million requests.'))

    def test_dates_are_not_mistaken_for_invented_metrics(self):
        from jobs.services.ats_checker import fabricated_metrics
        tailored = ORIGINAL_CV_NO_NUMBERS.replace(
            'January 2021 - Present', 'Jan 2021 - Present'
        )
        self.assertEqual(fabricated_metrics(ORIGINAL_CV_NO_NUMBERS, tailored), [])

    @override_settings(OPENAI_API_KEY='sk-test', ATS_TARGET_SCORE=90,
                       ATS_MAX_TAILOR_ATTEMPTS=3)
    @mock.patch('jobs.services.tailoring._retry_with_feedback')
    @mock.patch('jobs.services.tailoring.tailor_cv_for_job')
    def test_high_scoring_fabrication_loses_to_an_honest_lower_score(
        self, mock_tailor, mock_retry
    ):
        """The whole point: a dishonest draft never wins, however well it scores."""
        from jobs.services.tailoring import tailor_cv_for_job_with_ats

        # Draft 1: stuffed with invented skills and metrics -> scores well.
        dishonest = ATS_GOOD_CV + '\n- Delivered CI/CD pipelines, cutting costs 60%.\n'
        # Draft 2: honest, and necessarily scores lower.
        honest = ORIGINAL_CV_NO_NUMBERS
        mock_tailor.return_value = dishonest
        mock_retry.return_value = honest

        text, report, _attempts = tailor_cv_for_job_with_ats(
            ORIGINAL_CV_NO_NUMBERS, self.JD, 'Backend Engineer', 'Acme',
        )
        self.assertEqual(text, honest)
        self.assertTrue(report['honest'])
        self.assertFalse(report['unsupported_claims'])
        self.assertFalse(report['fabricated_metrics'])

    @override_settings(OPENAI_API_KEY='sk-test', ATS_MAX_TAILOR_ATTEMPTS=2)
    @mock.patch('jobs.services.tailoring._retry_with_feedback')
    @mock.patch('jobs.services.tailoring.tailor_cv_for_job')
    def test_persistent_fabrication_is_rejected_not_shipped(self, mock_tailor, mock_retry):
        """If every draft lies, ship the TRUTH — not the best-scoring lie.

        This test used to assert the opposite: that the fabrication was kept and
        merely flagged. That was the bug. Flagging a lie still sends the lie to
        the employer; the candidate is the one who pays for it at interview.
        """
        from jobs.services.tailoring import tailor_cv_for_job_with_ats
        liar = ORIGINAL_CV_NO_NUMBERS + '\n- Ran CI/CD, improving speed by 55%.\n'
        mock_tailor.return_value = liar
        mock_retry.return_value = liar

        text, report, _a = tailor_cv_for_job_with_ats(
            ORIGINAL_CV_NO_NUMBERS, self.JD, 'Backend Engineer', 'Acme',
        )
        # The untailored (true) CV is what ships.
        self.assertEqual(text, ORIGINAL_CV_NO_NUMBERS)
        self.assertTrue(report['fell_back_to_original'])
        # And the rejected fabrication is recorded, so the user can see why.
        rejected = ' '.join(report['fabrication_rejected'])
        self.assertIn('ci/cd', rejected)
        self.assertIn('55%', rejected)
        # What ships is honest, and carries the original's real score.
        self.assertTrue(report['honest'])
        self.assertEqual(report['unsupported_claims'], [])


class ATSCategoryBreakdownTests(TestCase):
    def test_report_exposes_five_weighted_categories(self):
        from jobs.services.ats_checker import check_cv_against_job
        report = check_cv_against_job(ATS_GOOD_CV, ATS_JOB_DESCRIPTION,
                                      'Senior Backend Engineer', 'London, UK')
        categories = report['categories']
        self.assertEqual(
            set(categories),
            {'keyword_matching', 'experience_relevance', 'formatting', 'education',
             'section_completeness'},
        )
        for cat in categories.values():
            self.assertIn('score', cat)
            self.assertIn('weight', cat)
            self.assertIn('weighted_score', cat)
            self.assertTrue(0 <= cat['score'] <= 100)
        # Keyword matching is the heaviest category, as in a real ATS.
        self.assertEqual(
            max(categories, key=lambda k: categories[k]['weight']), 'keyword_matching'
        )
        # The tailoring target, read from settings (80) — not a hardcoded number.
        self.assertEqual(report['score_needed'], settings.ATS_TARGET_SCORE)
        self.assertEqual(report['score_needed'], 80)


# A modern-stack advert: almost none of these tools are in the 115-word vocab.
MODERN_JD = """Senior Analytics Engineer - London

What We're Looking For
- Strong dbt and Snowflake experience is essential
- Build pipelines with Dagster and Fivetran
- Model data for Looker dashboards
- Terraform for infrastructure
- Excellent communication

Benefits
- 25 days holiday and private healthcare
"""

MODERN_CONTRACT = {
    'job_title': 'senior analytics engineer',
    'title_variants': ['analytics engineer', 'data engineer'],
    'hard_skills': ['dbt', 'snowflake', 'dagster', 'fivetran', 'looker', 'terraform'],
    'acronyms': [['ci/cd', 'continuous integration']],
    'soft_skills': ['communication'],
    'must_have': ['dbt', 'snowflake'],
    'source': 'test',
}


class JobKeywordContractTests(TestCase):
    def test_contract_helpers_flatten_both_acronym_forms(self):
        from jobs.services.job_keywords import all_contract_terms
        terms = all_contract_terms(MODERN_CONTRACT)
        self.assertIn('dbt', terms)
        self.assertIn('ci/cd', terms)
        self.assertIn('continuous integration', terms)

    def test_term_present_is_word_boundary_safe(self):
        from jobs.services.job_keywords import term_present
        self.assertTrue(term_present('ci/cd', 'we use ci/cd pipelines'))
        self.assertTrue(term_present('r', 'i code in r daily'))
        self.assertFalse(term_present('r', 'i code in ruby daily'))
        self.assertFalse(term_present('dbt', 'we use dbtx'))

    @override_settings(OPENAI_API_KEY='')
    def test_fallback_contract_without_openai_still_finds_skills(self):
        from jobs.services.job_keywords import extract_job_keywords
        contract = extract_job_keywords(MODERN_JD, 'Senior Analytics Engineer')
        self.assertEqual(contract['source'], 'fallback')
        self.assertIn('terraform', contract['hard_skills'])
        # The fallback cannot know what is mandatory, so it claims nothing rather
        # than guessing and knocking a candidate out on a guess.
        self.assertEqual(contract['must_have'], [])

    @override_settings(OPENAI_API_KEY='')
    def test_fallback_excludes_benefits_boilerplate(self):
        from jobs.services.job_keywords import extract_job_keywords
        contract = extract_job_keywords(MODERN_JD)
        for noise in ('holiday', 'healthcare', 'benefits'):
            self.assertNotIn(noise, contract['hard_skills'])

    @override_settings(OPENAI_API_KEY='sk-test')
    @mock.patch('openai.OpenAI')
    def test_openai_contract_is_normalised(self, mock_openai):
        from jobs.services.job_keywords import extract_job_keywords
        payload = {
            'job_title': 'Senior Analytics Engineer',
            'title_variants': ['Analytics Engineer'],
            'hard_skills': ['dbt', 'Snowflake', 'DBT', 'experience', 'x' * 200],
            'acronyms': [['CI/CD', 'Continuous Integration'], ['bad']],
            'soft_skills': ['Communication'],
            'must_have': ['dbt', 'not-a-real-skill'],
        }
        mock_openai.return_value.chat.completions.create.return_value = mock.Mock(
            choices=[mock.Mock(message=mock.Mock(content=json.dumps(payload)))]
        )
        contract = extract_job_keywords(MODERN_JD, 'Senior Analytics Engineer')

        self.assertEqual(contract['source'], 'openai')
        # Lowercased and de-duplicated; noise words and over-long phrases dropped.
        self.assertEqual(contract['hard_skills'], ['dbt', 'snowflake'])
        self.assertNotIn('experience', contract['hard_skills'])
        # must_have is only meaningful as a subset of hard_skills.
        self.assertEqual(contract['must_have'], ['dbt'])
        # Malformed acronym pairs are discarded, valid ones lowercased.
        self.assertEqual(contract['acronyms'], [['ci/cd', 'continuous integration']])

    @override_settings(OPENAI_API_KEY='sk-test')
    @mock.patch('openai.OpenAI')
    def test_openai_failure_falls_back_never_raises(self, mock_openai):
        from jobs.services.job_keywords import extract_job_keywords
        mock_openai.side_effect = RuntimeError('api down')
        contract = extract_job_keywords(MODERN_JD, 'Senior Analytics Engineer')
        self.assertEqual(contract['source'], 'fallback')
        self.assertTrue(contract['hard_skills'])

    def test_empty_description_returns_empty_contract(self):
        from jobs.services.job_keywords import extract_job_keywords
        contract = extract_job_keywords('', 'Engineer')
        self.assertEqual(contract['hard_skills'], [])
        self.assertEqual(contract['source'], 'empty')


class ContractScoringTests(TestCase):
    """The Jobscan-style number: coverage of the terms the job actually asked for."""

    def _cv(self, skills):
        return f"""JANE DOE
jane@example.com | London

PROFESSIONAL PROFILE
Senior Analytics Engineer.

KEY SKILLS
{chr(10).join('- ' + s for s in skills)}

PROFESSIONAL EXPERIENCE
Analytics Engineer | Acme | London
Jan 2019 - Present
- Did the work.

EDUCATION
BSc Computing | Leeds | 2015 - 2018
"""

    def test_full_coverage_scores_high(self):
        from jobs.services.ats_checker import score_cv_against_contract
        cv = self._cv(['dbt', 'Snowflake', 'Dagster', 'Fivetran', 'Looker',
                       'Terraform', 'CI/CD'])
        result = score_cv_against_contract(cv, MODERN_CONTRACT)
        self.assertGreaterEqual(result['score'], 90)
        self.assertEqual(result['missing_hard'], [])
        self.assertEqual(result['missing_must'], [])
        self.assertTrue(result['title_ok'])

    def test_missing_must_haves_are_reported_and_cost_the_score(self):
        from jobs.services.ats_checker import score_cv_against_contract
        cv = self._cv(['Looker', 'Terraform'])
        result = score_cv_against_contract(cv, MODERN_CONTRACT)
        self.assertIn('dbt', result['missing_must'])
        self.assertIn('snowflake', result['missing_must'])
        self.assertLess(result['score'], 60)

    def test_acronym_expansion_counts_as_coverage(self):
        """A CV saying 'continuous integration' satisfies an advert's 'ci/cd'."""
        from jobs.services.ats_checker import score_cv_against_contract
        cv = self._cv(['dbt', 'Snowflake', 'Dagster', 'Fivetran', 'Looker',
                       'Terraform', 'Continuous Integration'])
        result = score_cv_against_contract(cv, MODERN_CONTRACT)
        self.assertEqual(result['missing_acronyms'], [])

    def test_genuine_missing_terms_puts_must_haves_first(self):
        from jobs.services.ats_checker import (
            genuine_missing_terms,
            score_cv_against_contract,
        )
        result = score_cv_against_contract(self._cv(['Looker']), MODERN_CONTRACT)
        terms = genuine_missing_terms(result)
        self.assertEqual(set(terms[:2]), {'dbt', 'snowflake'})  # must-haves lead
        self.assertIn('dagster', terms)

    def test_absent_categories_do_not_award_free_marks(self):
        """A category with nothing to measure must not score 100.

        It used to: an empty must_have (25% weight) and no acronyms (10%) handed
        out 35 free points, which floored an unrelated job at ~45/100 against any
        well-formatted CV. The weight of an inapplicable category is now
        redistributed across the ones that actually apply.
        """
        from jobs.services.ats_checker import score_cv_against_contract
        unrelated = {
            'job_title': 'pastry chef',
            'title_variants': [],
            'hard_skills': ['baking', 'patisserie', 'cake decoration'],
            'acronyms': [],       # nothing to measure
            'soft_skills': [],
            'must_have': [],      # nothing to measure
        }
        result = score_cv_against_contract(ATS_GOOD_CV, unrelated)
        self.assertLess(result['score'], 25, 'unrelated job scored too high')
        self.assertIsNone(result['breakdown']['must_have'])
        self.assertIsNone(result['breakdown']['acronyms'])

    def test_scoring_never_raises_on_junk(self):
        from jobs.services.ats_checker import score_cv_against_contract
        # Nothing to score against -> None ("not scored"), never a number.
        self.assertIsNone(score_cv_against_contract('', MODERN_CONTRACT)['score'])
        self.assertIsNone(score_cv_against_contract('cv text', None)['score'])

    def test_an_empty_contract_is_not_scored_rather_than_scored_100(self):
        """The bug this replaces: an advert with no minable skills scored 100/100.

        With no hard skills, must-haves, acronyms or title, the only measurable
        category left was 'sections' — the CV's own heading coverage, which has
        nothing to do with the job. A well-formatted CV therefore scored a perfect
        100 against an advert nobody could read, while the reason column cheerfully
        said "no screenable skills could be mined from this job description".
        """
        from jobs.services.ats_checker import score_cv_against_contract
        empty_contract = {
            'job_title': '', 'title_variants': [], 'hard_skills': [],
            'acronyms': [], 'soft_skills': [], 'must_have': [], 'source': 'empty',
        }
        result = score_cv_against_contract(ATS_GOOD_CV, empty_contract)
        self.assertIsNone(result['score'])
        self.assertFalse(result['scorable'])

    def test_a_real_contract_is_scorable(self):
        from jobs.services.ats_checker import score_cv_against_contract
        result = score_cv_against_contract(ATS_GOOD_CV, MODERN_CONTRACT)
        self.assertTrue(result['scorable'])
        self.assertIsInstance(result['score'], int)

    def test_contract_catches_skills_the_hardcoded_vocab_cannot_see(self):
        """The whole point: dbt/Snowflake/Dagster are invisible to SKILL_VOCAB."""
        from jobs.services.keyword_extractor import extract_skills_from_text
        from jobs.services.ats_checker import score_cv_against_contract

        vocab_view = extract_skills_from_text(MODERN_JD)
        self.assertNotIn('snowflake', vocab_view)
        self.assertNotIn('dbt', vocab_view)

        # A CV with none of the real tools would look fine to the old vocabulary,
        # but the contract scores it for what it is.
        cv = self._cv(['Terraform', 'Communication'])
        self.assertLess(score_cv_against_contract(cv, MODERN_CONTRACT)['score'], 55)

    def test_fabrication_guard_covers_contract_skills(self):
        """An invented 'snowflake' must be caught, though no vocabulary knows it."""
        from jobs.services.ats_checker import unsupported_claims
        original = 'Jane Doe. I use Looker and Terraform.'
        tailored = original + ' Also expert in Snowflake and dbt.'
        claims = unsupported_claims(original, tailored, MODERN_JD, MODERN_CONTRACT)
        self.assertIn('snowflake', claims)
        self.assertIn('dbt', claims)
        self.assertNotIn('looker', claims)

    def test_compositional_terms_are_evidenced_word_by_word(self):
        """"dbt testing" is genuine if the CV says "using dbt, adding tests".

        Restating real experience in the advert's wording is the entire job of
        tailoring. A guard that flagged this would cry wolf on every honest
        rewrite and make the honesty signal worthless.
        """
        from jobs.services.ats_checker import unsupported_claims
        contract = dict(MODERN_CONTRACT,
                        hard_skills=['dbt testing', 'dbt macros', 'curated marts'],
                        must_have=[])
        original = 'Rebuilt warehouse models using dbt, adding tests and macros.'
        tailored = original + ' Led dbt testing and wrote dbt macros.'
        claims = unsupported_claims(original, tailored, MODERN_JD, contract)
        self.assertNotIn('dbt testing', claims)
        self.assertNotIn('dbt macros', claims)

    def test_compositional_check_still_catches_real_invention(self):
        from jobs.services.ats_checker import unsupported_claims
        contract = dict(MODERN_CONTRACT,
                        hard_skills=['curated marts', 'performance tuning'],
                        must_have=[])
        original = 'Rebuilt warehouse models using dbt, adding tests and macros.'
        # Neither "curated"/"marts" nor "performance"/"tuning" appear anywhere.
        tailored = original + ' Built curated marts and led performance tuning.'
        claims = unsupported_claims(original, tailored, MODERN_JD, contract)
        self.assertIn('curated marts', claims)
        self.assertIn('performance tuning', claims)

    def test_reworded_claims_are_flagged_for_review_not_blocked(self):
        """"cost optimisation" for "tuning warehouse costs" is a rewording.

        Blocking these would fire on nearly every honest rewrite, and a warning
        that always fires is a warning nobody reads. They are surfaced for the
        candidate to eyeball instead.
        """
        from jobs.services.ats_checker import (
            claims_needing_review,
            unsupported_claims,
        )
        contract = dict(MODERN_CONTRACT,
                        hard_skills=['cost optimisation', 'snowflake'],
                        must_have=[])
        original = 'Ran Snowflake and spent time tuning warehouse costs.'
        tailored = original + ' Led cost optimisation on Snowflake.'

        # Grounded in "costs", so not an invention -> does not block the draft.
        self.assertNotIn(
            'cost optimisation',
            unsupported_claims(original, tailored, MODERN_JD, contract),
        )
        # But surfaced, because "optimisation" is our word, not the candidate's.
        self.assertIn(
            'cost optimisation',
            claims_needing_review(original, tailored, MODERN_JD, contract),
        )

    def test_wholly_invented_single_word_skill_always_blocks(self):
        from jobs.services.ats_checker import (
            claims_needing_review,
            unsupported_claims,
        )
        contract = dict(MODERN_CONTRACT, hard_skills=['snowflake'], must_have=[])
        original = 'I use Looker and Terraform.'
        tailored = original + ' Expert in Snowflake.'
        self.assertIn(
            'snowflake', unsupported_claims(original, tailored, MODERN_JD, contract)
        )
        self.assertNotIn(
            'snowflake',
            claims_needing_review(original, tailored, MODERN_JD, contract),
        )


class ClaimRouterAdversarialTests(TestCase):
    """A fabricated multi-word skill must never be silently accepted.

    Word presence is strong evidence FOR grounding a reword and weak evidence
    AGAINST a fabrication; the two are not symmetric. A CV full of common words
    can supply every component of a skill the candidate never had. Proximity does
    not rescue this either — "ran risk reports ... and modelled churn" puts both
    words of "risk modelling" in one bullet. No lexical rule separates "did X"
    from "the words of X appear near each other", so compositional matches are
    surfaced for review rather than waved through.
    """

    # Every word of each fabricated skill below appears somewhere in this CV, in
    # an unrelated context. The candidate did none of them.
    CV = """Priya Nair
Data Engineer, Northwind Retail
- Loaded data from source systems and checked quality of the feeds each morning.
- Wrote tests for our ingestion scripts.
- Ran risk reports for the finance team and modelled churn for marketing.
- Managed access controls and reviewed security of the warehouse.
- Presented analysis to stakeholders; owned the incident response rota.
"""
    FAKES = [
        'data quality testing',
        'risk modelling',
        'security incident response',
        'source control management',
    ]

    def _route(self, terms):
        from jobs.services.ats_checker import (
            claims_needing_review,
            unsupported_claims,
        )
        contract = {'hard_skills': terms, 'must_have': [], 'acronyms': [],
                    'soft_skills': []}
        tailored = self.CV + '\n' + '\n'.join(f'- Led {t}.' for t in terms)
        return (
            unsupported_claims(self.CV, tailored, '', contract),
            claims_needing_review(self.CV, tailored, '', contract),
        )

    def test_no_fabricated_multiword_skill_is_silently_accepted(self):
        blocked, amber = self._route(self.FAKES)
        for fake in self.FAKES:
            self.assertIn(
                fake, blocked + amber,
                f'"{fake}" was silently accepted: every component word appears on '
                f'the CV, but the candidate never did it.',
            )

    def test_distinctive_invented_tool_still_blocks(self):
        blocked, amber = self._route(['snowflake', 'kubernetes'])
        self.assertIn('snowflake', blocked)
        self.assertIn('kubernetes', blocked)
        self.assertNotIn('snowflake', amber)

    def test_amber_claims_carry_the_line_that_grounds_them(self):
        from jobs.services.ats_checker import claim_evidence
        evidence = claim_evidence(self.CV, 'risk modelling')
        self.assertTrue(evidence)
        # The receipt is what makes the warning checkable rather than clickable-past.
        self.assertIn('risk reports', evidence[0]['line'])
        self.assertEqual(set(evidence[0]['matched']), {'risk', 'model'})

    def test_genuine_reword_still_lands_in_amber_not_block(self):
        """The behaviour we must not break while closing the leak."""
        from jobs.services.ats_checker import (
            claims_needing_review,
            unsupported_claims,
        )
        original = 'Ran Snowflake and spent time tuning warehouse costs.'
        tailored = original + ' Led cost optimisation on Snowflake.'
        contract = {'hard_skills': ['cost optimisation'], 'must_have': [],
                    'acronyms': [], 'soft_skills': []}
        self.assertNotIn(
            'cost optimisation',
            unsupported_claims(original, tailored, '', contract),
        )
        self.assertIn(
            'cost optimisation',
            claims_needing_review(original, tailored, '', contract),
        )


class PrescoreRecallTests(TestCase):
    """Stage-1 triage must not inherit the hardcoded vocabulary's blindness."""

    CV_SKILLS = ['dbt', 'snowflake', 'dagster', 'fivetran', 'looker', 'sql', 'python']
    RARE_STACK_JOB = ('Senior Analytics Engineer. You will use dbt, Snowflake, '
                      'Dagster, Fivetran and Looker daily.')
    GENERIC_JOB = ('Junior Developer. Some Python and SQL, plus Excel. '
                   'Agile team, good communication.')

    def test_perfect_rare_stack_match_outranks_a_mediocre_generic_job(self):
        from jobs.services.keyword_extractor import (
            extract_skills_from_text,
            prescore_job,
        )
        rare = prescore_job(
            self.CV_SKILLS, extract_skills_from_text(self.RARE_STACK_JOB),
            self.RARE_STACK_JOB,
        )
        generic = prescore_job(
            self.CV_SKILLS, extract_skills_from_text(self.GENERIC_JOB),
            self.GENERIC_JOB,
        )
        # The old vocab-only score saw NO skills in the rare-stack advert and gave
        # it a neutral 50, ranking it below vocab-rich jobs -- so the best match
        # could be cut before a contract was ever built for it.
        self.assertGreater(rare, generic)
        self.assertGreaterEqual(rare, 70)

    def test_direct_overlap_is_vocabulary_free(self):
        from jobs.services.keyword_extractor import (
            direct_overlap_score,
            extract_skills_from_text,
        )
        # SKILL_VOCAB cannot see any of this advert's tools...
        self.assertEqual(extract_skills_from_text(self.RARE_STACK_JOB), [])
        # ...but matching the CV's own LLM-derived skills against it needs no vocab.
        self.assertGreaterEqual(
            direct_overlap_score(self.CV_SKILLS, self.RARE_STACK_JOB), 70
        )

    def test_prescore_takes_the_best_signal_not_the_average(self):
        """Triage errs towards recall: one reason to look promising is enough."""
        from jobs.services.keyword_extractor import prescore_job
        # Vocab overlap is 0, direct overlap is high -> the job survives triage.
        self.assertGreaterEqual(
            prescore_job(self.CV_SKILLS, [], self.RARE_STACK_JOB), 70
        )

    def test_no_cv_skills_stays_neutral(self):
        from jobs.services.keyword_extractor import prescore_job
        # Nothing to match on -> neutral, so we never cut a job on no information.
        self.assertEqual(prescore_job([], [], self.RARE_STACK_JOB), 50)


class ATSFileChecksTests(TestCase):
    """Phase 1 checks that need the real file, not just its text."""

    def _write_docx(self, build):
        document = docx.Document()
        build(document)
        path = os.path.join(_TEST_MEDIA, f'ats_{id(build)}.docx')
        document.save(path)
        return path

    def test_docx_with_tables_and_images_is_flagged(self):
        from jobs.services.ats_checker import ATSChecker

        def build(document):
            document.add_paragraph('JOHN SMITH')
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = 'Skills'
            table.cell(0, 1).text = 'Python'

        path = self._write_docx(build)
        checker = ATSChecker(ATS_GOOD_CV, ATS_JOB_DESCRIPTION, file_path=path)
        result = checker.check_prohibited_elements()
        self.assertFalse(result['pass'])
        self.assertTrue(any('table' in e for e in result['elements']))

    def test_clean_docx_passes(self):
        from jobs.services.ats_checker import ATSChecker

        def build(document):
            for line in ATS_GOOD_CV.splitlines():
                document.add_paragraph(line)

        path = self._write_docx(build)
        checker = ATSChecker(ATS_GOOD_CV, ATS_JOB_DESCRIPTION, file_path=path)
        self.assertTrue(checker.check_prohibited_elements()['pass'])
        self.assertTrue(checker.check_layout()['pass'])

    def test_image_based_cv_fails_phase1(self):
        """A scanned CV yields almost no text — an ATS reads it as empty."""
        from jobs.services.ats_checker import check_cv_format

        def build(document):
            document.add_paragraph('CV')  # a scan yields little/no real text

        path = self._write_docx(build)
        result = check_cv_format('CV', file_path=path)
        self.assertFalse(result['pass'])
        self.assertFalse(result['file_format_ok'])
        self.assertIn('image-based', result['issue'])

    def test_missing_headings_fail_phase1(self):
        from jobs.services.ats_checker import check_cv_format
        result = check_cv_format(ATS_BAD_CV)
        self.assertFalse(result['pass'])
        self.assertTrue(result['missing_headers'])

    def test_text_only_check_skips_file_phases(self):
        """Without a file path the text phases still run; file checks are skipped."""
        from jobs.services.ats_checker import check_cv_format
        result = check_cv_format(ATS_GOOD_CV)
        self.assertTrue(result['pass'])
        self.assertTrue(result['file_checks_skipped'])


@override_settings(OPENAI_API_KEY='', MEDIA_ROOT=_TEST_MEDIA)
class ATSTailoringLoopTests(TestCase):
    def test_returns_report_without_openai(self):
        from jobs.services.tailoring import tailor_cv_for_job_with_ats
        text, report, attempts = tailor_cv_for_job_with_ats(
            ATS_GOOD_CV, ATS_JOB_DESCRIPTION, 'Senior Backend Engineer', 'Acme',
            'London, UK',
        )
        # No OpenAI -> the original CV is returned, scored honestly, no retries.
        self.assertEqual(text, ATS_GOOD_CV)
        self.assertEqual(attempts, 1)
        self.assertGreaterEqual(report['overall_score'], 75)

    @override_settings(OPENAI_API_KEY='sk-test', ATS_TARGET_SCORE=90,
                       ATS_MAX_TAILOR_ATTEMPTS=3)
    @mock.patch('jobs.services.tailoring._retry_with_feedback')
    @mock.patch('jobs.services.tailoring.tailor_cv_for_job')
    def test_retries_until_target_then_stops(self, mock_tailor, mock_retry):
        from jobs.services.tailoring import tailor_cv_for_job_with_ats
        # The source CV is the strong one, so a draft built from it is honest.
        # First draft is weak; the retry surfaces what was already there.
        mock_tailor.return_value = ATS_BAD_CV
        mock_retry.return_value = ATS_GOOD_CV

        text, report, attempts = tailor_cv_for_job_with_ats(
            ATS_GOOD_CV, ATS_JOB_DESCRIPTION, 'Senior Backend Engineer', 'Acme',
            'London, UK',
        )
        self.assertEqual(text, ATS_GOOD_CV)
        self.assertEqual(attempts, 2)  # stopped as soon as the target was cleared
        self.assertGreaterEqual(report['overall_score'], 75)
        self.assertTrue(report['honest'])
        mock_retry.assert_called_once()

    @override_settings(OPENAI_API_KEY='sk-test', ATS_TARGET_SCORE=100,
                       ATS_MAX_TAILOR_ATTEMPTS=2)
    @mock.patch('jobs.services.tailoring._retry_with_feedback')
    @mock.patch('jobs.services.tailoring.tailor_cv_for_job')
    def test_unreachable_target_keeps_best_draft_and_reports_true_score(
        self, mock_tailor, mock_retry
    ):
        from jobs.services.tailoring import tailor_cv_for_job_with_ats
        mock_tailor.return_value = ATS_BAD_CV
        mock_retry.return_value = ATS_GOOD_CV  # better, but still under 100

        text, report, attempts = tailor_cv_for_job_with_ats(
            ATS_GOOD_CV, ATS_JOB_DESCRIPTION, 'Senior Backend Engineer', 'Acme',
        )
        # Keeps the best attempt and reports its real score, not a flattering one.
        self.assertEqual(text, ATS_GOOD_CV)
        self.assertEqual(attempts, 2)
        self.assertLess(report['overall_score'], 100)


@override_settings(MEDIA_ROOT=_TEST_MEDIA, OPENAI_API_KEY='')
class ATSWorkflowIntegrationTests(TestCase):
    def setUp(self):
        from jobs.models import CV as CVModel
        self.user = User.objects.create_user(username='ats', password='pw12345!')
        self.cv = CVModel.objects.create(
            user=self.user, name='John Smith',
            original_file=SimpleUploadedFile('cv.docx', build_docx_bytes()),
            parsed_text=ATS_GOOD_CV,
            # Enough overlap with the job's skills to clear the stage-1 keyword
            # pre-screen, so the job actually reaches scoring and tailoring.
            parsed_data={
                'skills': ['python', 'django', 'sql', 'postgresql', 'aws', 'docker',
                           'kubernetes', 'rest', 'communication', 'leadership'],
                'job_titles': ['Backend Engineer'],
            },
        )
        self.run = SearchRun.objects.create(
            user=self.user, cv=self.cv, countries=['United Kingdom'],
            status=SearchRun.STATUS_PENDING,
        )

    def _raw_job(self, **overrides):
        job = {
            'title': 'Senior Backend Engineer', 'company': 'Acme',
            'location': 'London, UK', 'salary': '£80,000',
            'description': ATS_JOB_DESCRIPTION, 'applyLink': 'https://x/1',
            'datePosted': '', 'employmentType': '', 'seniorityLevel': '',
        }
        job.update(overrides)
        return job

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks._score_job')
    @mock.patch('jobs.tasks.search_jobs')
    def test_ats_report_generated_for_tailored_job(self, mock_search, mock_score, _sheets):
        from jobs.models import ATSReport
        from jobs.tasks import run_job_search
        mock_search.return_value = [self._raw_job()]
        mock_score.side_effect = fake_score(90, 'Strong match')

        run_job_search(self.run.pk)

        job = self.run.jobs.get()
        self.assertIsNotNone(job.ats_score)
        self.assertEqual(job.ats_status, Job.ATS_PASSED)
        report = ATSReport.objects.get(job=job)
        self.assertEqual(report.overall_score, job.ats_score)
        self.assertIn('phase3_keyword', report.report_data['phases'])

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks._score_job')
    @mock.patch('jobs.tasks.search_jobs')
    def test_former_knockouts_are_shown_scored_and_tailored(
        self, mock_search, mock_score, _sheets
    ):
        """The advert that used to bin the candidate now just gets scored.

        A PMP requirement, a degree requirement and a city the candidate does not
        live in were each a hard knock-out. Together they used to reject the job
        outright and skip tailoring. Now the job appears with its real score, and
        because that score clears the threshold it is tailored like any other.
        """
        from jobs.tasks import run_job_search
        mock_search.return_value = [self._raw_job(
            location='Aberdeen, UK',
            description=ATS_JOB_DESCRIPTION
            + '\nA valid PMP certification is required. Must hold a PhD. '
              '10 years of experience required. Based in Aberdeen.',
        )]
        mock_score.side_effect = fake_score(95, 'Strong match')

        run_job_search(self.run.pk)

        job = self.run.jobs.get()
        self.assertNotEqual(job.ats_status, 'REJECTED')
        self.assertEqual(job.match_score, 95)
        self.assertNotIn('Rejected', job.match_reason)
        self.assertTrue(job.tailored_pdf)

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks._score_job')
    @mock.patch('jobs.tasks.search_jobs')
    def test_badly_formatted_cv_does_not_reject_every_job(
        self, mock_search, mock_score, _sheets
    ):
        """A hard-to-parse CV is advice about the CV, not a verdict on the search."""
        from jobs.tasks import run_job_search
        self.cv.parsed_text = 'Python Django engineer, London UK. BSc Computer Science.'
        self.cv.save()
        mock_search.return_value = [self._raw_job()]
        mock_score.side_effect = fake_score(90, 'Strong match')

        run_job_search(self.run.pk)

        job = self.run.jobs.get()
        self.assertNotEqual(job.ats_status, 'REJECTED')
        self.assertTrue(job.tailored_pdf)  # still tailored despite poor formatting

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks._score_job')
    @mock.patch('jobs.tasks.search_jobs')
    def test_duplicate_application_to_same_company_is_flagged(
        self, mock_search, mock_score, _sheets
    ):
        """Phase 7: two CVs sent to the same company are flagged, as an ATS would."""
        from jobs.models import ATSReport
        from jobs.tasks import run_job_search
        # Two different roles at the same company, in one run.
        mock_search.return_value = [
            self._raw_job(title='Senior Backend Engineer'),
            self._raw_job(title='Backend Engineer', applyLink='https://x/2'),
        ]
        mock_score.side_effect = fake_score(90, 'Strong match')

        run_job_search(self.run.pk)

        reports = {r.job.title: r.report_data for r in ATSReport.objects.all()}
        self.assertEqual(len(reports), 2)
        # The second application to Acme is the one flagged as a duplicate.
        flagged = [t for t, d in reports.items() if d.get('duplicate_application')]
        self.assertEqual(len(flagged), 1)

    @override_settings(ATS_THRESHOLD=100)
    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks._score_job')
    @mock.patch('jobs.tasks.search_jobs')
    def test_unreachable_ats_threshold_only_flags(self, mock_search, mock_score, _s):
        """An unreachable ATS threshold flags the job. It cannot reject it.

        There is no strict mode any more: a threshold is a label on a score, and
        the only status a low one can produce is BELOW_THRESHOLD.
        """
        from jobs.tasks import run_job_search
        mock_search.return_value = [self._raw_job()]
        mock_score.side_effect = fake_score(90, 'Strong match')

        run_job_search(self.run.pk)
        job = self.run.jobs.get()
        self.assertEqual(job.ats_status, Job.ATS_BELOW_THRESHOLD)
        self.assertTrue(job.tailored_pdf)
        self.assertNotIn('REJECTED', dict(Job.ATS_STATUS_CHOICES))


@override_settings(MEDIA_ROOT=_TEST_MEDIA, OPENAI_API_KEY='')
class ATSReportViewTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username='viewer', password='pw12345!')
        self.client.login(username='viewer', password='pw12345!')
        self.run = SearchRun.objects.create(user=self.user)
        self.job = Job.objects.create(
            search_run=self.run, title='Backend Engineer', company='Acme',
            location='London', application_link='https://x/1', ats_score=88,
            ats_status=Job.ATS_PASSED,
            # Results only show jobs at or above the match threshold.
            match_score=90,
        )

    def _make_report(self):
        from jobs.models import ATSReport
        from jobs.services.ats_checker import check_cv_against_job
        data = check_cv_against_job(ATS_GOOD_CV, ATS_JOB_DESCRIPTION,
                                    'Senior Backend Engineer', 'London, UK')
        return ATSReport.objects.create(
            job=self.job, overall_score=data['overall_score'], report_data=data,
        )

    def test_report_page_renders_all_phases(self):
        self._make_report()
        response = self.client.get(reverse('ats_report', args=[self.job.pk]))
        self.assertEqual(response.status_code, 200)
        for needle in ['Phase 1', 'Phase 3', 'ATS score']:
            self.assertContains(response, needle)
        # The knock-out phase is gone from the report, and so is its verdict.
        self.assertNotContains(response, 'Knock-out')
        self.assertNotContains(response, 'ATS Rejected')

    def test_missing_report_redirects(self):
        response = self.client.get(reverse('ats_report', args=[self.job.pk]))
        self.assertRedirects(response, reverse('search_results', args=[self.run.pk]))

    def test_report_scoped_to_owner(self):
        other = User.objects.create_user(username='snoop', password='pw12345!')
        run = SearchRun.objects.create(user=other)
        job = Job.objects.create(search_run=run, title='X', company='Y',
                                 location='Z', application_link='https://x/2')
        response = self.client.get(reverse('ats_report', args=[job.pk]))
        self.assertEqual(response.status_code, 404)

    def test_results_page_shows_ats_column(self):
        self._make_report()
        response = self.client.get(reverse('search_results', args=[self.run.pk]))
        self.assertContains(response, '88')
        self.assertContains(response, reverse('ats_report', args=[self.job.pk]))


class CityFilterTests(TestCase):
    def test_cities_are_listed_per_country(self):
        from jobs.services.locations import cities_for_country
        self.assertIn('Bristol', cities_for_country('United Kingdom'))
        self.assertIn('Bristol', cities_for_country('uk'))  # code or name
        self.assertIn('Berlin', cities_for_country('Germany'))
        self.assertEqual(cities_for_country('Narnia'), [])

    def test_city_map_is_keyed_by_the_names_the_form_offers(self):
        from jobs.services.locations import cities_by_country_name
        mapping = cities_by_country_name()
        self.assertIn('United Kingdom', mapping)
        self.assertIn('London', mapping['United Kingdom'])

    def test_city_goes_into_the_actor_input(self):
        from jobs.services.apify_service import _build_actor_input
        inp = _build_actor_input('United Kingdom', 30000, 50, city='Bristol')
        # The actor searches the city, not just the country.
        self.assertEqual(inp['custom_location'], 'Bristol, United Kingdom')
        self.assertEqual(inp['country'], 'uk')
        # No city -> country-level search, as before.
        self.assertEqual(
            _build_actor_input('United Kingdom', None, 50)['custom_location'],
            'United Kingdom',
        )

    def test_location_filter_keeps_only_the_chosen_city(self):
        from jobs.services.apify_service import _filter_by_location
        jobs = [
            {'title': 'A', 'location': 'Bristol, UK'},
            {'title': 'B', 'location': 'London, UK'},
            {'title': 'C', 'location': 'Greater Bristol'},
            {'title': 'D', 'location': ''},  # unknown -> kept, not dropped
        ]
        kept = {j['title'] for j in _filter_by_location(jobs, ['United Kingdom'],
                                                        city='Bristol')}
        self.assertEqual(kept, {'A', 'C', 'D'})

    def test_no_city_falls_back_to_country_matching(self):
        from jobs.services.apify_service import _filter_by_location
        jobs = [{'title': 'A', 'location': 'London, UK'},
                {'title': 'B', 'location': 'Berlin, Germany'}]
        kept = {j['title'] for j in _filter_by_location(jobs, ['United Kingdom'])}
        self.assertEqual(kept, {'A'})

    def test_empty_city_result_is_not_backfilled(self):
        """If the user asked for Bristol, don't hand them a page of London jobs."""
        from jobs.services.apify_service import _filter_by_location
        jobs = [{'title': 'A', 'location': 'London, UK'}]
        self.assertEqual(_filter_by_location(jobs, ['United Kingdom'], city='Bristol'), [])

    def test_preferences_form_rejects_a_city_outside_the_chosen_country(self):
        from jobs.forms import UserPreferencesForm
        form = UserPreferencesForm(data={
            'target_countries': ['United Kingdom'], 'target_city': 'Berlin',
            'currency': 'GBP',
        })
        self.assertFalse(form.is_valid())
        self.assertIn('target_city', form.errors)

    def test_preferences_form_accepts_a_valid_city(self):
        from jobs.forms import UserPreferencesForm
        form = UserPreferencesForm(data={
            'target_countries': ['United Kingdom'], 'target_city': 'Bristol',
            'currency': 'GBP',
        })
        self.assertTrue(form.is_valid(), form.errors)
        self.assertEqual(form.cleaned_data['target_city'], 'Bristol')


class CityPersistenceTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username='cityuser', password='pw12345!')
        self.client.login(username='cityuser', password='pw12345!')

    def test_city_saved_and_threaded_into_the_search_run(self):
        response = self.client.post(reverse('edit_preferences'), {
            'target_countries': ['United Kingdom'], 'target_city': 'Manchester',
            'salary_min': '40000', 'currency': 'GBP',
        })
        self.assertEqual(response.status_code, 302)
        prefs = UserPreferences.objects.get(user=self.user)
        self.assertEqual(prefs.target_city, 'Manchester')

        _make_cv_for(self.user)
        with mock.patch('jobs.views.process_job_search.delay'):
            self.client.post(reverse('start_search'))
        run = SearchRun.objects.get(user=self.user)
        self.assertEqual(run.city, 'Manchester')  # snapshot on the run


@override_settings(MEDIA_ROOT=_TEST_MEDIA, OPENAI_API_KEY='')
class MatchThresholdDisplayTests(TestCase):
    """Sub-75 jobs are stored but shown nowhere — not in the app, not in Excel."""

    def setUp(self):
        self.user = User.objects.create_user(username='thr', password='pw12345!')
        self.client.login(username='thr', password='pw12345!')
        self.run = SearchRun.objects.create(
            user=self.user, status=SearchRun.STATUS_COMPLETED,
        )
        self.good = Job.objects.create(
            search_run=self.run, title='Strong Match', company='Acme',
            location='London', match_score=88, application_link='https://x/1',
        )
        self.weak = Job.objects.create(
            search_run=self.run, title='Weak Match', company='Beta',
            location='Leeds', match_score=40, application_link='https://x/2',
        )

    def test_results_page_shows_every_job_across_three_tabs(self):
        """The results page shows ALL jobs; the tabs filter, they don't hide."""
        response = self.client.get(reverse('search_results', args=[self.run.pk]))
        self.assertContains(response, 'Strong Match')
        self.assertContains(response, 'Weak Match')
        self.assertEqual(response.context['total_found'], 2)
        self.assertEqual(response.context['above_count'], 1)
        self.assertEqual(response.context['below_count'], 1)

    def test_results_page_never_says_rejected(self):
        """No job is rejected any more, so the word cannot appear on the page."""
        Job.objects.create(
            search_run=self.run, title='Unreadable Advert', company='Gamma',
            location='Cardiff', match_score=None, application_link='https://x/3',
        )
        response = self.client.get(reverse('search_results', args=[self.run.pk]))
        content = response.content.decode()
        self.assertNotIn('Rejected', content)
        self.assertNotIn('ATS Rejected', content)

    def test_unscored_job_is_shown_last_and_labelled(self):
        unscored = Job.objects.create(
            search_run=self.run, title='Unreadable Advert', company='Gamma',
            location='Cardiff', match_score=None, application_link='https://x/3',
        )
        response = self.client.get(reverse('search_results', args=[self.run.pk]))

        # Sorted last: unknown is not the same as zero, and it must not outrank
        # the job that scored 40.
        self.assertEqual(
            [j.pk for j in response.context['jobs']],
            [self.good.pk, self.weak.pk, unscored.pk],
        )
        self.assertEqual(response.context['not_scored_count'], 1)
        self.assertContains(response, 'Not scored')
        # It is never counted as "above 75".
        self.assertEqual(response.context['above_count'], 1)
        self.assertEqual(response.context['total_found'], 3)

    def test_every_row_carries_the_score_the_tabs_filter_on(self):
        response = self.client.get(reverse('search_results', args=[self.run.pk]))
        content = response.content.decode()
        # The tabs filter client-side on data-score, so it must be on every row.
        self.assertIn('data-score="88"', content)
        self.assertIn('data-score="40"', content)
        for tab in ('data-filter="all"', 'data-filter="above"', 'data-filter="below"'):
            self.assertIn(tab, content)

    def test_weak_jobs_are_still_stored_not_deleted(self):
        self.assertEqual(Job.objects.filter(search_run=self.run).count(), 2)
        self.assertTrue(Job.objects.filter(pk=self.weak.pk).exists())

    def test_excel_export_excludes_sub_threshold_jobs(self):
        from openpyxl import load_workbook
        response = self.client.get(reverse('export_excel', args=[self.run.pk]))
        wb = load_workbook(io.BytesIO(b''.join(response.streaming_content)))
        titles = [row[0] for row in wb['Jobs'].iter_rows(min_row=2, values_only=True)]
        self.assertIn('Strong Match', titles)
        self.assertNotIn('Weak Match', titles)


@override_settings(MEDIA_ROOT=_TEST_MEDIA, OPENAI_API_KEY='')
class SearchProgressPhaseTests(TestCase):
    """Progress must climb through its phases, not sit at 0 and snap to 100."""

    def setUp(self):
        self.user = User.objects.create_user(username='prog', password='pw12345!')
        _make_cv_for(self.user)
        self.run = SearchRun.objects.create(
            user=self.user, countries=['United Kingdom'], min_salary=30000,
            status=SearchRun.STATUS_PENDING,
        )

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks._score_job')
    @mock.patch('jobs.tasks.search_jobs')
    def test_progress_moves_through_the_phases(self, mock_search, mock_score, _sheets):
        from jobs.tasks import run_job_search
        mock_search.return_value = TWO_RAW_JOBS
        mock_score.side_effect = fake_score(85, 'Strong match')

        seen = []
        original = SearchRun.save

        def record(self_run, *args, **kwargs):
            if self_run.pk == self.run.pk:
                seen.append(self_run.progress)
            return original(self_run, *args, **kwargs)

        with mock.patch.object(SearchRun, 'save', record):
            run_job_search(self.run.pk)

        self.assertTrue(seen)
        # Monotonic, and it visits the intermediate bands rather than jumping.
        self.assertEqual(seen, sorted(seen))
        self.assertIn(15, seen)  # fetch complete
        self.assertTrue(any(15 < p < 75 for p in seen), f'no scoring phase: {seen}')
        self.assertTrue(any(75 <= p < 100 for p in seen), f'no tailoring phase: {seen}')
        self.assertEqual(seen[-1], 100)

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks._score_job')
    @mock.patch('jobs.tasks.search_jobs')
    def test_parallel_scoring_preserves_per_job_results(self, mock_search, mock_score, _s):
        """Concurrency must not shuffle which score belongs to which job."""
        from jobs.tasks import run_job_search

        mock_search.return_value = TWO_RAW_JOBS
        # Distinct score per job, so a mix-up across threads would be visible.
        mock_score.side_effect = fake_score(by_description={
            'Django': (90, 'backend'),
            'Entry role': (62, 'entry'),
        })
        run_job_search(self.run.pk)

        jobs = {j.title: j for j in self.run.jobs.all()}
        self.assertEqual(jobs['Backend Engineer'].match_score, 90)
        self.assertEqual(jobs['Backend Engineer'].match_reason, 'backend')
        # Each job keeps its own real score and its own real reason — no salary
        # note bolted on the front, because salary no longer rules anything out.
        self.assertEqual(jobs['Junior Dev'].match_score, 62)
        self.assertEqual(jobs['Junior Dev'].match_reason, 'entry')


@override_settings(MEDIA_ROOT=_TEST_MEDIA, OPENAI_API_KEY='')
class ScoreVarianceTests(TestCase):
    """Scores must be real and per-job, not a constant."""

    def setUp(self):
        from jobs.models import CV as CVModel
        self.user = User.objects.create_user(username='var', password='pw12345!')
        self.cv = CVModel.objects.create(
            user=self.user, name='Priya',
            original_file=SimpleUploadedFile('cv.docx', build_docx_bytes()),
            parsed_text=ATS_GOOD_CV,
            parsed_data={'skills': ['python', 'django', 'aws', 'docker'],
                         'job_titles': ['Backend Engineer']},
        )
        self.run = SearchRun.objects.create(
            user=self.user, cv=self.cv, countries=['United Kingdom'],
            status=SearchRun.STATUS_PENDING,
        )

    def _raw(self, title, description):
        return {
            'title': title, 'company': 'Acme', 'location': 'London',
            'salary': '£70,000', 'description': description,
            'applyLink': 'https://x/1', 'datePosted': '',
            'employmentType': '', 'seniorityLevel': '',
        }

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks.search_jobs')
    def test_scores_vary_with_the_job_description(self, mock_search, _sheets):
        """Real scoring, no mocking of the scorer: different JDs, different scores."""
        from jobs.tasks import run_job_search
        mock_search.return_value = [
            # Close to the CV (Python/Django/AWS/Docker/PostgreSQL).
            self._raw('Backend Engineer', ATS_JOB_DESCRIPTION),
            # Nothing to do with it.
            self._raw('Pastry Chef', 'We need a pastry chef. Baking, patisserie, '
                                     'cake decoration and food hygiene essential.'),
            # Partial overlap.
            self._raw('Data Analyst', 'Python and SQL required. Tableau, Excel, '
                                      'statistics and data analysis.'),
        ]

        run_job_search(self.run.pk)

        scores = {j.title: j.match_score for j in self.run.jobs.all()}
        # Every job has a real score — none is a placeholder 0.
        self.assertEqual(len(scores), 3)
        self.assertTrue(all(s is not None for s in scores.values()))
        # And they genuinely differ: the aligned job beats the unrelated one.
        self.assertGreater(scores['Backend Engineer'], scores['Pastry Chef'])
        self.assertGreater(len(set(scores.values())), 1, f'constant scores: {scores}')

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks.search_jobs')
    def test_score_measures_the_original_cv_not_the_tailored_one(
        self, mock_search, _sheets
    ):
        """match_score answers "does this candidate already fit?"

        It must be computed against the ORIGINAL CV. The tailored CV's coverage is
        a different question, and it lives in ats_score.
        """
        from jobs.services.ats_checker import score_cv_against_contract
        from jobs.services.job_keywords import extract_job_keywords
        from jobs.tasks import run_job_search

        mock_search.return_value = [self._raw('Backend Engineer', ATS_JOB_DESCRIPTION)]
        run_job_search(self.run.pk)

        job = self.run.jobs.get()
        expected = score_cv_against_contract(
            self.cv.parsed_text,
            extract_job_keywords(ATS_JOB_DESCRIPTION, 'Backend Engineer'),
        )['score']
        self.assertEqual(job.match_score, expected)

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks.search_jobs')
    def test_an_empty_description_is_not_scored_rather_than_scored_100(
        self, mock_search, _sheets
    ):
        """The headline bug: an unreadable advert used to score a perfect 100.

        With no keywords to mine, the only measurable thing left was the CV's own
        section coverage — so a well-formatted CV scored 100/100 against an empty
        advert. It must come back as "not scored" (None) instead, and sort below
        every job that has a real score.
        """
        from jobs.tasks import run_job_search
        mock_search.return_value = [
            self._raw('Backend Engineer', ATS_JOB_DESCRIPTION),
            self._raw('Mystery Role', ''),
        ]

        run_job_search(self.run.pk)

        jobs = {j.title: j for j in self.run.jobs.all()}
        self.assertIsNone(jobs['Mystery Role'].match_score)
        self.assertIn('Not scored', jobs['Mystery Role'].match_reason)
        # It is still in the results — just unscored, and never tailored.
        self.assertEqual(len(jobs), 2)
        self.assertFalse(jobs['Mystery Role'].tailored_pdf)
        self.assertIsNotNone(jobs['Backend Engineer'].match_score)

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks.search_jobs')
    def test_no_job_is_ever_scored_a_blanket_100(self, mock_search, _sheets):
        """No job may score 100 unless it genuinely covers everything the ad asks."""
        from jobs.tasks import run_job_search
        mock_search.return_value = [
            self._raw('Pastry Chef', 'Baking, patisserie and cake decoration.'),
            self._raw('Mystery Role', ''),
            self._raw('Welder', 'MIG welding, TIG welding, fabrication, blueprints.'),
        ]

        run_job_search(self.run.pk)

        for job in self.run.jobs.all():
            self.assertNotEqual(job.match_score, 100, f'{job.title} scored a fake 100')

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks._score_job')
    @mock.patch('jobs.tasks.search_jobs')
    def test_jobs_from_every_uk_city_reach_the_results(
        self, mock_search, mock_score, _sheets
    ):
        """No location filter survives: a CV in London does not lose Manchester.

        The location knock-out compared the advert's city to the CV's, so a
        candidate whose CV said "London" was rejected from every job outside it —
        and a CV with no readable address was rejected from all of them.
        """
        from jobs.tasks import run_job_search
        cities = ['London', 'Manchester', 'Edinburgh', 'Bristol', 'Cardiff',
                  'Leeds', 'Belfast', 'Aberdeen']
        mock_search.return_value = [
            dict(self._raw(f'Backend Engineer {city}', ATS_JOB_DESCRIPTION),
                 location=f'{city}, UK')
            for city in cities
        ]
        mock_score.side_effect = fake_score(80, 'Strong match')

        run_job_search(self.run.pk)

        jobs = list(self.run.jobs.all())
        self.assertEqual(len(jobs), len(cities))
        self.assertEqual(
            {j.location.split(',')[0] for j in jobs}, set(cities),
        )
        for job in jobs:
            self.assertNotEqual(job.ats_status, 'REJECTED')
            self.assertEqual(job.match_score, 80)


@override_settings(MEDIA_ROOT=_TEST_MEDIA, OPENAI_API_KEY='')
class SalaryPreferenceTests(TestCase):
    """Salary is the one optional filter, and by default it only flags."""

    def setUp(self):
        from jobs.models import CV as CVModel
        self.user = User.objects.create_user(username='sal', password='pw12345!')
        self.cv = CVModel.objects.create(
            user=self.user, name='Sam',
            original_file=SimpleUploadedFile('cv.docx', build_docx_bytes()),
            parsed_text=ATS_GOOD_CV,
            parsed_data={'skills': ['python', 'django'],
                         'job_titles': ['Backend Engineer']},
        )
        self.run = SearchRun.objects.create(
            user=self.user, cv=self.cv, countries=['United Kingdom'],
            max_salary=70000, status=SearchRun.STATUS_PENDING,
        )

    def _raw(self, title, salary):
        return {
            'title': title, 'company': 'Acme', 'location': 'London',
            'salary': salary, 'description': ATS_JOB_DESCRIPTION,
            'applyLink': 'https://x/1', 'datePosted': '',
            'employmentType': '', 'seniorityLevel': '',
        }

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks._score_job')
    @mock.patch('jobs.tasks.search_jobs')
    def test_over_ceiling_job_is_flagged_not_hidden(
        self, mock_search, mock_score, _sheets
    ):
        from jobs.tasks import run_job_search
        mock_search.return_value = [
            self._raw('Within budget', '£60,000'),
            self._raw('Over budget', '£95,000'),
        ]
        mock_score.side_effect = fake_score(85, 'Strong match')

        run_job_search(self.run.pk)

        jobs = {j.title: j for j in self.run.jobs.all()}
        self.assertEqual(len(jobs), 2)  # nothing dropped
        self.assertTrue(jobs['Over budget'].above_salary_preference)
        self.assertIn('Above your salary preference', jobs['Over budget'].match_reason)
        # Flagged, but still worth pursuing: the tailored CV is still generated.
        self.assertTrue(jobs['Over budget'].tailored_pdf)
        self.assertFalse(jobs['Within budget'].above_salary_preference)

    @override_settings(SALARY_HARD_FILTER=True)
    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks._score_job')
    @mock.patch('jobs.tasks.search_jobs')
    def test_hard_filter_drops_only_when_explicitly_enabled(
        self, mock_search, mock_score, _sheets
    ):
        from jobs.tasks import run_job_search
        mock_search.return_value = [
            self._raw('Within budget', '£60,000'),
            self._raw('Over budget', '£95,000'),
        ]
        mock_score.side_effect = fake_score(85, 'Strong match')

        run_job_search(self.run.pk)

        titles = {j.title for j in self.run.jobs.all()}
        self.assertEqual(titles, {'Within budget'})

    @mock.patch('jobs.tasks.GoogleSheetsLogger')
    @mock.patch('jobs.tasks._score_job')
    @mock.patch('jobs.tasks.search_jobs')
    def test_unparseable_salary_is_never_flagged(self, mock_search, mock_score, _s):
        """"Competitive" is not evidence of anything, least of all overpayment."""
        from jobs.tasks import run_job_search
        mock_search.return_value = [self._raw('Mystery pay', 'Competitive')]
        mock_score.side_effect = fake_score(85, 'Strong match')

        run_job_search(self.run.pk)

        job = self.run.jobs.get()
        self.assertFalse(job.above_salary_preference)
        self.assertNotIn('Above your salary preference', job.match_reason)


def altered_facts_for(original, tailored):
    from jobs.services.ats_checker import altered_facts
    return altered_facts(original, tailored)


class FactIntegrityTests(TestCase):
    """Education and employment are facts. Tailoring may not touch them."""

    ORIGINAL = """JANE DOE
PROFESSIONAL EXPERIENCE
Analyst | Beta Corp | Leeds
Jan 2019 - Dec 2021

EDUCATION
BSc Mathematics | University of Leeds | 2015 - 2018
Lower Second Class (2:2)
"""

    def test_upgraded_grade_is_caught(self):
        from jobs.services.ats_checker import altered_facts
        tailored = self.ORIGINAL.replace('Lower Second Class (2:2)',
                                         'Upper Second Class (2:1)')
        problems = altered_facts(self.ORIGINAL, tailored)
        self.assertTrue(problems)
        self.assertTrue(any('2:1' in p and 'invented' in p for p in problems))
        self.assertTrue(any('2:2' in p and 'removed' in p for p in problems))

    def test_invented_degree_is_caught(self):
        from jobs.services.ats_checker import altered_facts
        tailored = self.ORIGINAL + '\nMSc Data Science | Imperial College | 2019\n'
        problems = altered_facts(self.ORIGINAL, tailored)
        self.assertTrue(any("Master's" in p for p in problems))

    def test_swapped_university_is_caught(self):
        from jobs.services.ats_checker import altered_facts
        tailored = self.ORIGINAL.replace('University of Leeds',
                                         'University of Cambridge')
        problems = altered_facts(self.ORIGINAL, tailored)
        self.assertTrue(any('cambridge' in p.lower() for p in problems))
        self.assertTrue(any('leeds' in p.lower() and 'removed' in p for p in problems))

    def test_moved_date_is_caught(self):
        from jobs.services.ats_checker import altered_facts
        tailored = self.ORIGINAL.replace('Jan 2019', 'Jan 2017')  # closing a gap
        problems = altered_facts(self.ORIGINAL, tailored)
        self.assertTrue(any('2017' in p for p in problems))

    def test_deleted_role_is_caught(self):
        """Dropping a real role is a misrepresentation by omission."""
        from jobs.services.ats_checker import altered_facts
        tailored = self.ORIGINAL.replace('Jan 2019 - Dec 2021', '')
        problems = altered_facts(self.ORIGINAL, tailored)
        self.assertTrue(any('removed' in p for p in problems))

    def test_honest_rewording_is_not_flagged(self):
        """Wording, emphasis and ordering are free. Only the facts are fixed."""
        from jobs.services.ats_checker import altered_facts
        tailored = self.ORIGINAL + (
            '\nKEY SKILLS\n- Python\n- Data Analysis\n'
            '\nPROFESSIONAL PROFILE\nAnalyst with Python and data analysis '
            'experience, seeking a Data Analyst role.\n'
        )
        # Every employer, title, date, degree and grade is untouched; the rewrite
        # only adds surfacing and re-emphasis. Nothing to flag.
        self.assertEqual(altered_facts(self.ORIGINAL, tailored), [])

    def test_retitled_role_is_caught(self):
        """"Analyst" -> "Senior Data Analyst" is a lie a recruiter will check.

        The most plausible-looking fabrication a tailoring model can make, and
        previously the one thing the guard let through — the prompt forbade it,
        but nothing verified it.
        """
        from jobs.services.ats_checker import altered_facts
        tailored = self.ORIGINAL.replace(
            'Analyst | Beta Corp | Leeds', 'Senior Data Analyst | Beta Corp | Leeds',
        )
        problems = altered_facts(self.ORIGINAL, tailored)
        self.assertTrue(
            any('job title' in p and 'invented' in p for p in problems), problems
        )

    def test_invented_employer_is_caught(self):
        from jobs.services.ats_checker import altered_facts
        tailored = self.ORIGINAL + (
            '\nEngineer | Google | London\nJan 2022 - Dec 2023\n'
        )
        problems = altered_facts(self.ORIGINAL, tailored)
        self.assertTrue(
            any('employer' in p and 'google' in p.lower() for p in problems), problems
        )

    def test_dropped_employer_is_caught(self):
        """A real role deleted to tidy the CV is a misrepresentation by omission."""
        from jobs.services.ats_checker import altered_facts
        two_roles = self.ORIGINAL + (
            '\nJunior Analyst | Gamma Ltd | Leeds\nJan 2017 - Dec 2018\n'
        )
        # The rewrite silently drops the older role.
        problems = altered_facts(two_roles, self.ORIGINAL)
        self.assertTrue(
            any('gamma' in p.lower() and 'removed' in p for p in problems), problems
        )

    def test_education_lines_are_not_mistaken_for_employers(self):
        """"BSc Maths | University of Leeds" uses the same pipe format as a role."""
        from jobs.services.ats_checker import employment_entries
        employers = {e for _t, e in employment_entries(self.ORIGINAL)}
        self.assertIn('beta corp', employers)
        self.assertNotIn('university of leeds', employers)

    @override_settings(OPENAI_API_KEY='sk-test', ATS_MAX_TAILOR_ATTEMPTS=2,
                       ATS_TARGET_SCORE=80)
    @mock.patch('jobs.services.tailoring._retry_with_feedback')
    @mock.patch('jobs.services.tailoring.tailor_cv_for_job')
    def test_honest_ceiling_below_target_is_kept_and_flagged(
        self, mock_tailor, mock_retry
    ):
        """Below 80 with no fabrication = the candidate's true ceiling for this job.

        The score is kept as-is and flagged. It is never topped up by claiming a
        skill the CV cannot evidence.
        """
        from jobs.services.tailoring import tailor_cv_for_job_with_ats
        # An honest draft: same facts, no invented skills, but a weak match for a
        # job wanting Kubernetes/Terraform the candidate has never touched.
        honest = self.ORIGINAL + '\nKEY SKILLS\n- Python\n'
        mock_tailor.return_value = honest
        mock_retry.return_value = honest

        text, report, _attempts = tailor_cv_for_job_with_ats(
            self.ORIGINAL,
            'We need Kubernetes, Terraform, Go and Kafka experience.',
            'Platform Engineer', 'Acme',
        )

        self.assertEqual(altered_facts_for(self.ORIGINAL, text), [])
        self.assertLess(report['overall_score'], 80)
        self.assertTrue(report['below_target_honestly'])
        self.assertEqual(report['honest_ceiling'], report['overall_score'])
        self.assertIn('highest honest score', report['ceiling_reason'])
        # Nothing was fabricated to close the gap.
        self.assertEqual(report['unsupported_claims'], [])
        self.assertTrue(report['honest'])

    @override_settings(OPENAI_API_KEY='sk-test', ATS_MAX_TAILOR_ATTEMPTS=2)
    @mock.patch('jobs.services.tailoring._retry_with_feedback')
    @mock.patch('jobs.services.tailoring.tailor_cv_for_job')
    def test_a_draft_that_alters_facts_never_ships(self, mock_tailor, mock_retry):
        """If every draft misstates the record, fall back to the untailored CV.

        A CV that lies about a degree is worse than one that scores badly.
        """
        from jobs.services.tailoring import tailor_cv_for_job_with_ats
        liar = self.ORIGINAL.replace('Lower Second Class (2:2)', 'First Class')
        mock_tailor.return_value = liar
        mock_retry.return_value = liar

        text, report, _attempts = tailor_cv_for_job_with_ats(
            self.ORIGINAL, 'Python role', 'Analyst', 'Acme',
        )
        self.assertEqual(text, self.ORIGINAL)  # fell back to the truth
        self.assertTrue(report['fell_back_to_original'])
        self.assertTrue(report['altered_facts_rejected'])

    @override_settings(OPENAI_API_KEY='sk-test', ATS_MAX_TAILOR_ATTEMPTS=2,
                       ATS_TARGET_SCORE=80)
    @mock.patch('jobs.services.tailoring._retry_with_feedback')
    @mock.patch('jobs.services.tailoring.tailor_cv_for_job')
    def test_a_fabricated_high_score_never_ships(self, mock_tailor, mock_retry):
        """A CV that invents skills must not ship, however well it scores.

        The loop used to keep the best-scoring dishonest draft and only log a
        warning. Driving it for real, a backend engineer's CV tailored to a
        Salesforce architect role came back claiming six Salesforce skills the
        candidate had never touched — and scored a perfect 100. It shipped. A
        document that wins the ATS screen and collapses at interview is worse than
        a low score, so there is now no path that returns a fabrication.
        """
        from jobs.services.tailoring import tailor_cv_for_job_with_ats
        # An honest CV, plus skills that appear nowhere in the original.
        liar = self.ORIGINAL + (
            '\nKEY SKILLS\n- Salesforce Apex\n- Kubernetes\n- Terraform\n'
        )
        mock_tailor.return_value = liar
        mock_retry.return_value = liar

        text, report, _attempts = tailor_cv_for_job_with_ats(
            self.ORIGINAL,
            'We need Salesforce Apex, Kubernetes and Terraform.',
            'Platform Engineer', 'Acme',
        )

        self.assertEqual(text, self.ORIGINAL)  # the truth shipped, not the lie
        self.assertTrue(report['fell_back_to_original'])
        self.assertTrue(report['fabrication_rejected'])
        # The reported score is the ORIGINAL CV's real one, not the fabrication's.
        self.assertTrue(report['honest'])
        self.assertEqual(report['unsupported_claims'], [])
        self.assertTrue(report['below_target_honestly'])
        self.assertLess(report['overall_score'], 80)


@override_settings(MEDIA_ROOT=_TEST_MEDIA)
class EtaAndTimestampTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username='eta', password='pw12345!')
        self.client.login(username='eta', password='pw12345!')

    def _run(self, **kwargs):
        return SearchRun.objects.create(user=self.user, **kwargs)

    def test_eta_extrapolates_from_work_done(self):
        from django.utils import timezone
        run = self._run(
            status=SearchRun.STATUS_RUNNING, progress=25,
            started_at=timezone.now() - timedelta(seconds=60),
        )
        # 60s bought 25% -> ~240s total -> ~180s remaining.
        eta = run.eta_seconds()
        self.assertIsNotNone(eta)
        self.assertAlmostEqual(eta, 180, delta=10)

    def test_no_eta_before_there_is_anything_to_extrapolate_from(self):
        from django.utils import timezone
        run = self._run(
            status=SearchRun.STATUS_RUNNING, progress=2,
            started_at=timezone.now() - timedelta(seconds=5),
        )
        # Too early to guess: a confident wrong number is worse than "estimating".
        self.assertIsNone(run.eta_seconds())

    def test_no_eta_when_not_running(self):
        from django.utils import timezone
        run = self._run(status=SearchRun.STATUS_COMPLETED, progress=100,
                        started_at=timezone.now())
        self.assertIsNone(run.eta_seconds())

    def test_status_endpoint_exposes_live_progress_phase_and_eta(self):
        from django.utils import timezone
        run = self._run(
            status=SearchRun.STATUS_RUNNING, progress=40, total_jobs=10,
            started_at=timezone.now() - timedelta(seconds=40),
        )
        data = self.client.get(reverse('search_status', args=[run.pk])).json()
        self.assertEqual(data['progress'], 40)
        self.assertEqual(data['phase'], 'Scoring against your CV')
        self.assertIsNotNone(data['eta_seconds'])
        self.assertIn('minute', data['eta_display'])

    def test_eta_display_says_estimating_when_unknown(self):
        run = self._run(status=SearchRun.STATUS_RUNNING, progress=1)
        data = self.client.get(reverse('search_status', args=[run.pk])).json()
        self.assertIsNone(data['eta_seconds'])
        self.assertEqual(data['eta_display'], 'Estimating…')

    def test_search_run_timestamp_renders_in_local_time_12_hour(self):
        """The list showed UTC in 24-hour format, so BST times were an hour out."""
        import zoneinfo
        from django.utils import timezone

        # The Search Runs table only renders once the user has a CV profile.
        cv = _make_cv_for(self.user)
        # 15:45 UTC on a summer's day = 4:45 PM in London (BST).
        when = datetime(2026, 7, 12, 15, 45, tzinfo=zoneinfo.ZoneInfo('UTC'))
        run = self._run(status=SearchRun.STATUS_COMPLETED, cv=cv)
        SearchRun.objects.filter(pk=run.pk).update(created_at=when)

        with override_settings(TIME_ZONE='Europe/London'):
            timezone.activate('Europe/London')
            try:
                response = self.client.get(reverse('dashboard'))
            finally:
                timezone.deactivate()

        content = response.content.decode()
        self.assertIn('12 Jul 2026, 04:45 PM', content)
        self.assertNotIn('15:45', content)  # not raw UTC, not 24-hour

    def test_timestamp_follows_the_configured_timezone(self):
        """TIME_ZONE is what makes the clock right, and it is read from .env.

        The template was always correct; leaving TIME_ZONE unset (so it fell back
        to Europe/London) is what made the Search Runs clock read hours out for a
        user elsewhere. Same instant, two zones, two correct renderings.
        """
        import zoneinfo
        from django.utils import timezone

        cv = _make_cv_for(self.user)
        # 15:45 UTC = 4:45 PM London = 8:45 PM Karachi (UTC+5).
        when = datetime(2026, 7, 12, 15, 45, tzinfo=zoneinfo.ZoneInfo('UTC'))
        run = self._run(status=SearchRun.STATUS_COMPLETED, cv=cv)
        SearchRun.objects.filter(pk=run.pk).update(created_at=when)

        with override_settings(TIME_ZONE='Asia/Karachi'):
            timezone.activate('Asia/Karachi')
            try:
                content = self.client.get(reverse('dashboard')).content.decode()
            finally:
                timezone.deactivate()

        self.assertIn('12 Jul 2026, 08:45 PM', content)
        self.assertNotIn('04:45 PM', content)


@override_settings(MEDIA_ROOT=_TEST_MEDIA)
class ProgressPollingWiringTests(TestCase):
    """The bar not moving is almost always a broken link in the chain.

    Endpoint URL, JSON field names and the JS selectors have to agree. Nothing
    else in the suite fails if one of them drifts — the page just quietly stops
    updating — so the contract between them is pinned here.
    """

    def setUp(self):
        self.user = User.objects.create_user(username='poll', password='pw12345!')
        self.client.login(username='poll', password='pw12345!')
        self.cv = _make_cv_for(self.user)
        self.run = SearchRun.objects.create(
            user=self.user, cv=self.cv, status=SearchRun.STATUS_RUNNING,
            progress=42, total_jobs=10,
            started_at=django_now() - timedelta(seconds=60),
        )

    def test_dashboard_renders_every_hook_the_polling_js_queries(self):
        content = self.client.get(reverse('dashboard')).content.decode()
        status_url = reverse('search_status', args=[self.run.pk])

        # The JS reads the URL off the row, then writes into these three nodes.
        self.assertIn(f'data-status-url="{status_url}"', content)
        self.assertIn(f'data-run-id="{self.run.pk}"', content)
        for selector in ('run-status', 'progress-bar', 'progress-text',
                         'progress-eta'):
            self.assertIn(selector, content)
        # The bar must be painted with the live value, not a hardcoded zero.
        self.assertIn('width:42%', content)

    def test_status_endpoint_returns_every_field_the_js_reads(self):
        data = self.client.get(
            reverse('search_status', args=[self.run.pk])
        ).json()
        for field in ('status', 'status_display', 'progress', 'phase',
                      'eta_seconds', 'eta_display', 'processed', 'total'):
            self.assertIn(field, data, f'JS reads data.{field}; endpoint omits it')
        self.assertEqual(data['progress'], 42)

    def test_bar_is_only_rendered_while_there_is_something_to_show(self):
        self.run.status = SearchRun.STATUS_COMPLETED
        self.run.save(update_fields=['status'])
        content = self.client.get(reverse('dashboard')).content.decode()
        # The bar's markup is gone (the class name still appears in the polling
        # script, which is why this looks for the element, not the string).
        self.assertNotIn('<div class="progress">', content)


class ProfilePageTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username='pia', password='pw12345!', email='pia@example.com',
        )

    def test_profile_requires_login(self):
        response = self.client.get(reverse('profile'))
        self.assertEqual(response.status_code, 302)
        self.assertIn(reverse('login'), response.url)

    def test_profile_shows_account_details(self):
        self.client.login(username='pia', password='pw12345!')
        UserPreferences.objects.create(
            user=self.user, target_countries=['United Kingdom'],
            target_city='Bristol', salary_min=50000, currency='GBP',
        )
        response = self.client.get(reverse('profile'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'pia')
        self.assertContains(response, 'pia@example.com')
        self.assertContains(response, 'Bristol')
        self.assertContains(response, 'Change password')

    def test_profile_never_shows_a_password(self):
        """Only a salted hash exists; there is nothing to show and we must not try."""
        self.client.login(username='pia', password='pw12345!')
        response = self.client.get(reverse('profile'))
        self.assertNotContains(response, 'pw12345!')
        self.assertNotContains(response, self.user.password)  # not even the hash

    def test_password_change_works(self):
        self.client.login(username='pia', password='pw12345!')
        response = self.client.post(reverse('password_change'), {
            'old_password': 'pw12345!',
            'new_password1': 'BrandNewPw99!',
            'new_password2': 'BrandNewPw99!',
        })
        self.assertRedirects(response, reverse('password_change_done'))
        self.user.refresh_from_db()
        self.assertTrue(self.user.check_password('BrandNewPw99!'))


@override_settings(MEDIA_ROOT=_TEST_MEDIA)
class ExcelExportTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username='frank', password='pw12345!')
        self.client.login(username='frank', password='pw12345!')

    def _run_with_job(self, status=SearchRun.STATUS_COMPLETED, score=85):
        run = SearchRun.objects.create(user=self.user, status=status)
        Job.objects.create(
            search_run=run, title='Backend Engineer', company='Acme',
            location='London', salary='£60,000', match_score=score,
            match_reason='Strong match', application_link='https://x.com/1',
        )
        return run

    def test_export_returns_xlsx(self):
        run = self._run_with_job()
        response = self.client.get(reverse('export_excel', args=[run.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        self.assertIn(f'search_results_{run.pk}.xlsx', response['Content-Disposition'])

        # Load the returned bytes back into openpyxl and verify content.
        from openpyxl import load_workbook
        content = b''.join(response.streaming_content)
        wb = load_workbook(io.BytesIO(content))
        ws = wb['Jobs']
        self.assertEqual(ws['A1'].value, 'Job Title')
        self.assertEqual(ws['I1'].value, 'Match Score')
        self.assertEqual(ws['A2'].value, 'Backend Engineer')
        self.assertEqual(ws['I2'].value, 85)
        self.assertEqual(ws.freeze_panes, 'A2')

    def test_export_blocked_when_not_completed(self):
        run = self._run_with_job(status=SearchRun.STATUS_RUNNING)
        response = self.client.get(reverse('export_excel', args=[run.pk]))
        self.assertEqual(response.status_code, 302)  # redirected with error


class GoogleSheetsTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username='gina', password='pw12345!')
        self.run = SearchRun.objects.create(user=self.user)
        self.job = Job.objects.create(
            search_run=self.run, title='Dev', company='Acme', location='London',
            match_score=90, application_link='https://x.com/1',
            job_skills=['python', 'sql'], missing_skills=['sql'], ats_score=92,
        )

    def test_disabled_when_not_configured(self):
        from jobs.services.google_sheets import GoogleSheetsLogger
        with override_settings(GOOGLE_SHEET_ID='', GOOGLE_SHEETS_CREDENTIALS_JSON=''):
            sheets = GoogleSheetsLogger()
        self.assertFalse(sheets.enabled)
        self.assertFalse(sheets.log_job(self.job, 'Gina'))

    def test_sanitize_tab_name(self):
        from jobs.services.google_sheets import sanitize_tab_name
        self.assertEqual(sanitize_tab_name('Haseeb Ijaz'), 'Haseeb Ijaz')
        self.assertEqual(sanitize_tab_name('A/B:C[D]'), 'A-B-C-D-')
        self.assertEqual(sanitize_tab_name(''), 'Candidate')
        self.assertEqual(len(sanitize_tab_name('x' * 80)), 50)

    def test_build_row_matches_headers(self):
        from jobs.services.google_sheets import HEADERS, GoogleSheetsLogger
        with override_settings(GOOGLE_SHEET_ID='', GOOGLE_SHEETS_CREDENTIALS_JSON=''):
            sheets = GoogleSheetsLogger()
        row = sheets.build_row(self.job, cv_skills=['python', 'django'])
        self.assertEqual(len(row), len(HEADERS))
        self.assertIn('Dev', row)
        self.assertEqual(row[HEADERS.index('Match Score')], 90)
        self.assertEqual(row[HEADERS.index('ATS Score')], 92)
        self.assertEqual(row[HEADERS.index('CV Parsed Skills')], 'python, django')
        self.assertEqual(row[HEADERS.index('Job Required Skills')], 'python, sql')
        self.assertEqual(row[HEADERS.index('Missing Skills')], 'sql')

    @mock.patch('jobs.services.google_sheets.os.path.exists', return_value=True)
    def test_creates_tab_per_candidate_and_appends(self, _exists):
        from jobs.services.google_sheets import GoogleSheetsLogger, HEADERS
        import gspread

        with override_settings(GOOGLE_SHEET_ID='sheet123',
                               GOOGLE_SHEETS_CREDENTIALS_JSON='/fake/creds.json'):
            with mock.patch('gspread.authorize') as authorize, \
                 mock.patch('google.oauth2.service_account.Credentials.from_service_account_file'):
                sheet = authorize.return_value.open_by_key.return_value
                # Candidate has no tab yet -> a new one is created with headers.
                sheet.worksheet.side_effect = gspread.exceptions.WorksheetNotFound('nope')
                new_tab = sheet.add_worksheet.return_value

                sheets = GoogleSheetsLogger()
                self.assertTrue(sheets.enabled)
                self.assertTrue(sheets.log_job(self.job, 'Haseeb Ijaz', cv_skills=['python']))

        sheet.add_worksheet.assert_called_once()
        self.assertEqual(sheet.add_worksheet.call_args.kwargs['title'], 'Haseeb Ijaz')
        # First append writes headers, second writes the job row.
        first_call = new_tab.append_row.call_args_list[0].args[0]
        second_call = new_tab.append_row.call_args_list[1].args[0]
        self.assertEqual(first_call, HEADERS)
        self.assertIn('Dev', second_call)
