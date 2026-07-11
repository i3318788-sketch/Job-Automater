"""Helpers for extracting text from uploaded CV files and resolving the active CV."""
import os

from django.core.exceptions import ValidationError

ALLOWED_EXTENSIONS = {'.pdf', '.docx'}

SESSION_ACTIVE_CV = 'active_cv_id'


def resolve_active_cv(request, profiles=None):
    """Return the user's active CV profile, honouring ?cv_id= then the session.

    Selecting via ?cv_id= persists the choice in the session. Falls back to the
    first profile, or None if the user has no CVs. ``profiles`` may be passed to
    avoid a duplicate query.
    """
    from .models import CV  # local import to avoid circulars

    if not request.user.is_authenticated:
        return None
    if profiles is None:
        profiles = list(CV.objects.filter(user=request.user).order_by('id'))
    if not profiles:
        request.session.pop(SESSION_ACTIVE_CV, None)
        return None

    by_id = {c.pk: c for c in profiles}

    # 1) Explicit selection via query param.
    raw = request.GET.get('cv_id')
    if raw and raw.isdigit() and int(raw) in by_id:
        request.session[SESSION_ACTIVE_CV] = int(raw)
        return by_id[int(raw)]

    # 2) Previously selected (and still valid).
    session_id = request.session.get(SESSION_ACTIVE_CV)
    if session_id in by_id:
        return by_id[session_id]

    # 3) Default to the first profile.
    request.session[SESSION_ACTIVE_CV] = profiles[0].pk
    return profiles[0]


def validate_cv_extension(filename):
    """Raise ValidationError if the file is not a PDF or DOCX."""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValidationError(
            'Unsupported file type "%(ext)s". Please upload a PDF or DOCX file.',
            params={'ext': ext or '(none)'},
        )
    return ext


def extract_text_from_pdf(file_obj):
    """Extract text from a PDF file-like object using PyPDF2."""
    from PyPDF2 import PdfReader

    reader = PdfReader(file_obj)
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or '')
        except Exception:
            # Skip pages that fail to extract rather than aborting the upload.
            continue
    return '\n'.join(parts).strip()


def extract_text_from_docx(file_obj):
    """Extract text from a DOCX file-like object using python-docx."""
    import docx

    document = docx.Document(file_obj)
    paragraphs = [p.text for p in document.paragraphs]
    # Include table cell text too, since CVs often use tables for layout.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return '\n'.join(paragraphs).strip()


def extract_cv_text(uploaded_file):
    """Return extracted text for an uploaded PDF or DOCX file."""
    ext = validate_cv_extension(uploaded_file.name)
    # Rewind in case the file has been read during validation.
    uploaded_file.seek(0)
    if ext == '.pdf':
        return extract_text_from_pdf(uploaded_file)
    return extract_text_from_docx(uploaded_file)
