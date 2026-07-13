"""ATS-style CV checker.

Simulates how an Applicant Tracking System reads and ranks a CV against a job
description, in seven phases:

    1. Parsing & technical formatting  (hard filter — needs the CV *file*)
    2. Knock-out filters               (hard filter — needs the job's requirements)
    3. Keyword matching                (scored)
    4. Context & proximity             (scored)
    5. Experience & chronology         (scored)
    6. Education parsing               (scored)
    7. Final scoring & ranking

Phases 1 and 2 are pass/fail: failing either means the CV would never reach a
human, so the caller should treat the result as "ATS Rejected" regardless of the
score. Phases 3-6 produce the weighted 0-100 score.

The checker is deterministic and offline — no API calls — so it is cheap enough
to run on every job in a search, and its scores are reproducible.

``pdfplumber`` is used for the file-level checks (columns, images, tables,
fonts) when installed. Without it those checks are reported as ``skipped``
rather than failed, and the text-based phases still run.
"""
import hashlib
import logging
import os
import re
from datetime import date

from django.conf import settings

from .keyword_extractor import SKILL_VOCAB
from .pdf_generator import SECTION_ALIASES, parse_cv_sections

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Tunables
# ---------------------------------------------------------------------------
# Weight of each scored phase in the overall score. Keyword matching dominates,
# as it does in a real ATS.
PHASE_WEIGHTS = {
    'phase1_parsing': 0.10,
    'phase3_keyword': 0.45,
    'phase4_context': 0.20,
    'phase5_experience': 0.15,
    'phase6_education': 0.10,
}

# A CV's minimum defensible score. Overridable via ATS_THRESHOLD in .env.
DEFAULT_THRESHOLD = 75

# Sections an ATS expects to find. Missing any of these is a parsing failure.
CRITICAL_SECTIONS = {'experience', 'education', 'skills'}

# Fonts that every ATS parser renders reliably.
STANDARD_FONTS = {
    'arial', 'calibri', 'times new roman', 'times', 'verdana', 'helvetica',
    'georgia', 'garamond', 'tahoma', 'cambria', 'book antiqua', 'trebuchet ms',
    'liberation serif', 'liberation sans', 'dejavu sans', 'roboto', 'lato',
}

# Soft skills carry less weight than hard skills, exactly as in a real ATS.
SOFT_SKILLS = [
    'communication', 'leadership', 'teamwork', 'collaboration', 'problem solving',
    'problem-solving', 'time management', 'adaptability', 'creativity',
    'critical thinking', 'attention to detail', 'interpersonal', 'negotiation',
    'presentation', 'organisation', 'organization', 'mentoring', 'coaching',
    'stakeholder management', 'analytical', 'self-motivated', 'proactive',
]

# Hard skills = the shared vocabulary, minus anything that is really a soft skill.
_SOFT_SET = {s.lower() for s in SOFT_SKILLS}
HARD_SKILL_VOCAB = [s for s in SKILL_VOCAB if s.lower() not in _SOFT_SET]

# Domain synonyms. This replaces WordNet, which knows nothing about "Kubernetes"
# or "SEMrush" but would happily equate "lead" with "conduce". Keys and values
# are matched case-insensitively; the relation is treated as symmetric.
SYNONYMS = {
    'javascript': ['js', 'ecmascript'],
    'typescript': ['ts'],
    'python': ['py'],
    'postgresql': ['postgres', 'psql'],
    'mongodb': ['mongo'],
    'kubernetes': ['k8s'],
    'amazon web services': ['aws'],
    'google cloud platform': ['gcp', 'google cloud'],
    'microsoft azure': ['azure'],
    'machine learning': ['ml'],
    'artificial intelligence': ['ai'],
    'natural language processing': ['nlp'],
    'continuous integration': ['ci', 'ci/cd'],
    'continuous deployment': ['cd', 'ci/cd'],
    'search engine optimisation': ['seo', 'search engine optimization'],
    'pay per click': ['ppc', 'paid search'],
    'customer relationship management': ['crm'],
    'user experience': ['ux'],
    'user interface': ['ui'],
    'quality assurance': ['qa', 'testing'],
    'project management': ['programme management', 'program management'],
    'business intelligence': ['bi'],
    'extract transform load': ['etl'],
    'representational state transfer': ['rest', 'restful'],
    'application programming interface': ['api', 'apis'],
    'structured query language': ['sql'],
    'power bi': ['powerbi'],
    'microsoft excel': ['excel', 'ms excel'],
    'node.js': ['node', 'nodejs'],
    'next.js': ['nextjs'],
    'react': ['react.js', 'reactjs'],
    'vue': ['vue.js', 'vuejs'],
    '.net': ['dotnet', 'asp.net'],
    'c#': ['csharp'],
    'agile': ['scrum', 'kanban'],
    'devops': ['sre', 'site reliability'],
}

# Certifications an ATS knock-out filter commonly searches for by exact name.
CERTIFICATIONS = [
    'pmp', 'prince2', 'capm', 'csm', 'safe', 'cpa', 'acca', 'cima', 'cfa',
    'shrm-cp', 'shrm-scp', 'cipd', 'phr', 'sphr', 'cissp', 'cisa', 'cism',
    'comptia security+', 'comptia a+', 'comptia network+', 'ccna', 'ccnp',
    'aws certified solutions architect', 'aws certified developer',
    'aws certified', 'azure fundamentals', 'az-900', 'microsoft certified',
    'google analytics individual qualification', 'google ads certification',
    'itil', 'six sigma', 'lean six sigma', 'togaf', 'ceh', 'oscp',
    'chartered engineer', 'rics', 'acca', 'aat',
]

# Degree hierarchy. Higher number = higher level; a knock-out compares levels.
DEGREE_LEVELS = [
    (5, 'Doctorate', [
        'phd', 'ph.d', 'doctorate', 'doctoral', 'dphil', 'md', 'edd',
    ]),
    (4, "Master's", [
        'master', 'masters', "master's", 'msc', 'm.sc', 'ma', 'm.a', 'mba',
        'meng', 'm.eng', 'mres', 'mphil', 'llm', 'mst', 'ms',
    ]),
    (3, "Bachelor's", [
        'bachelor', 'bachelors', "bachelor's", 'bsc', 'b.sc', 'ba', 'b.a',
        'beng', 'b.eng', 'bcom', 'llb', 'bs', 'undergraduate degree',
    ]),
    (2, 'Foundation/Associate', [
        'associate degree', 'foundation degree', 'hnd', 'hnc', 'fdsc', 'fda',
    ]),
    (1, 'Diploma/Certificate', [
        'diploma', 'certificate', 'a-level', 'a levels', 'btec', 'nvq', 'gcse',
    ]),
]

# UK degree classifications mapped onto a 4.0-style scale for comparison.
UK_CLASSIFICATIONS = {
    'first class': 4.0, 'first-class': 4.0, '1st class': 4.0,
    'upper second': 3.5, '2:1': 3.5, '2.1': 3.5, 'upper second class': 3.5,
    'lower second': 3.0, '2:2': 3.0, '2.2': 3.0, 'lower second class': 3.0,
    'third class': 2.5, '3rd class': 2.5,
    'distinction': 4.0, 'merit': 3.3, 'pass': 2.5,
}

WORK_AUTH_PATTERNS = [
    ('Citizen', r'\b(?:uk|us|british|american|eu)\s+citizen(?:ship)?\b|\bcitizen of\b'),
    ('Permanent Resident', r'\b(?:green card|permanent resident|indefinite leave to remain|ilr|settled status)\b'),
    ('Authorised to Work', r'\b(?:right to work|authoriz?sed to work|authorized to work|work permit|no sponsorship required|eligible to work)\b'),
    ('Requires Sponsorship', r'\b(?:requires? sponsorship|need(?:s|ing)? sponsorship|visa sponsorship required|tier 2 sponsorship|skilled worker visa)\b'),
]

RELOCATION_RE = re.compile(
    r'\b(?:willing to relocate|open to relocation|happy to relocate|will relocate|'
    r'relocating to|open to relocating)\b',
    re.IGNORECASE,
)

# "5+ years", "at least 3 years", "minimum of 7 years' experience"
YEARS_REQUIRED_RE = re.compile(
    r'(\d{1,2})\s*\+?\s*(?:-\s*\d{1,2}\s*)?(?:or more\s+)?year[s]?(?:’|\')?'
    r'(?:\s+of)?(?:\s+(?:relevant|proven|professional|commercial|hands-on|industry))?'
    r'(?:\s+work)?\s+experience',
    re.IGNORECASE,
)

# Numbers, percentages, currency, and magnitudes — evidence of quantified impact.
QUANTIFIER_RE = re.compile(
    r'(?:\d+(?:[.,]\d+)*\s*%|[£$€]\s?\d|\b\d+(?:[.,]\d+)*\s*(?:k|m|bn|million|billion)\b|\b\d+(?:[.,]\d+)*\b)',
    re.IGNORECASE,
)

GPA_RE = re.compile(r'\bgpa[:\s]*(?:of\s*)?(\d\.\d{1,2})\b', re.IGNORECASE)

# "Minimum GPA 3.5", "GPA of 3.5 or higher", "minimum 2:1 degree"
GPA_REQUIRED_RE = re.compile(
    r'(?:minimum|min\.?|at least|above|requires?)\s*(?:a\s*)?(?:gpa\s*(?:of\s*)?)?'
    r'(\d\.\d{1,2})|gpa\s*(?:of\s*)?(\d\.\d{1,2})\s*(?:or\s+(?:higher|above)|\+)',
    re.IGNORECASE,
)

MONTHS = {
    'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6, 'jul': 7,
    'aug': 8, 'sep': 9, 'sept': 9, 'oct': 10, 'nov': 11, 'dec': 12,
}

# "Jan 2022 - Present", "2017 – 2020", "01/2019 to 12/2020"
DATE_RANGE_RE = re.compile(
    r'(?P<start>(?:[A-Za-z]{3,9}\.?\s+)?(?:\d{1,2}/)?\d{4})'
    r'\s*(?:[-–—]|to|until)\s*'
    r'(?P<end>present|current|now|ongoing|(?:[A-Za-z]{3,9}\.?\s+)?(?:\d{1,2}/)?\d{4})',
    re.IGNORECASE,
)

BULLET_PREFIXES = ('•', '-', '*', '–', '—', '·', 'o ')

# How many tokens apart the words of a multi-word requirement may sit and still
# count as "used together" in one bullet point.
PROXIMITY_WINDOW = 10

STOPWORDS = {
    'a', 'an', 'the', 'and', 'or', 'but', 'if', 'then', 'else', 'for', 'to', 'of',
    'in', 'on', 'at', 'by', 'with', 'from', 'as', 'is', 'are', 'was', 'were', 'be',
    'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would',
    'shall', 'should', 'can', 'could', 'may', 'might', 'must', 'this', 'that',
    'these', 'those', 'it', 'its', 'we', 'our', 'you', 'your', 'they', 'their',
    'i', 'me', 'my', 'he', 'she', 'his', 'her', 'us', 'them', 'who', 'whom',
    'which', 'what', 'when', 'where', 'why', 'how', 'all', 'any', 'both', 'each',
    'more', 'most', 'other', 'some', 'such', 'no', 'not', 'only', 'own', 'same',
    'so', 'than', 'too', 'very', 'just', 'also', 'about', 'into', 'over', 'up',
    'out', 'down', 'across', 'within', 'role', 'job', 'work', 'working', 'team',
    'company', 'business', 'candidate', 'candidates', 'applicant', 'applicants',
    'you\'ll', 'we\'re', 'etc', 'e.g', 'i.e', 'per', 'via', 'new', 'good', 'great',
    'strong', 'excellent', 'ability', 'able', 'looking', 'seeking', 'join',
    'opportunity', 'apply', 'please', 'benefits', 'salary', 'position', 'well',
}


# ---------------------------------------------------------------------------
# Text normalisation, stemming, synonyms
# ---------------------------------------------------------------------------
# NLTK's PorterStemmer is used when NLTK happens to be installed (it needs no
# corpus download); otherwise the suffix stripper below does the job. Either way
# both the CV and the job description are stemmed with the *same* function, so
# matching stays consistent.
try:  # pragma: no cover - depends on optional dependency
    from nltk.stem import PorterStemmer as _PorterStemmer

    _porter = _PorterStemmer()
except Exception:  # pragma: no cover
    _porter = None

# WordNet needs a corpus download (`python -m nltk.downloader wordnet`), so it is
# strictly optional: without it the curated synonym map below still applies.
try:  # pragma: no cover - depends on optional corpus
    from nltk.corpus import wordnet as _wordnet

    _wordnet.synsets('test')  # force the lookup now, not mid-search
except Exception:  # pragma: no cover
    _wordnet = None

_SUFFIXES = (
    'ization', 'isation', 'ements', 'ement', 'ments', 'ment', 'ities', 'ility',
    'ively', 'ingly', 'edly', 'ness', 'ions', 'ion', 'ing', 'ies', 'ied', 'ers',
    'er', 'ed', 'es', 's', 'ly', 'ance', 'ence', 'al', 'ive',
)


def _stem(word):
    """Reduce a word to a matching key (manage/managing/managed/management -> manag).

    Short tokens (<= 3 chars) are returned untouched so acronyms like SQL, AWS
    and QA survive intact.
    """
    w = word.lower()
    if len(w) <= 3:
        return w
    if _porter is not None:  # pragma: no cover - optional dependency
        return _porter.stem(w)

    # Two passes, so "communications" -> "communication" -> "communicat".
    for _ in range(2):
        for suffix in _SUFFIXES:
            if w.endswith(suffix) and len(w) - len(suffix) >= 3:
                w = w[: -len(suffix)]
                break
        else:
            break
    # "planned" -> "plann" -> "plan"
    if len(w) > 3 and w[-1] == w[-2] and w[-1] not in 'aeiou':
        w = w[:-1]
    # "manage" -> "manag", so it matches the stem of "managing".
    if len(w) > 3 and w.endswith('e'):
        w = w[:-1]
    return w


def _stem_phrase(phrase):
    """Stem every word of a phrase into a single comparable key."""
    return ' '.join(_stem(t) for t in re.findall(r"[a-z0-9+#./]+", phrase.lower()))


def _build_synonym_index():
    """Flatten SYNONYMS into term -> set(equivalent terms), symmetrically."""
    index = {}
    for canonical, alternatives in SYNONYMS.items():
        group = {canonical.lower(), *(a.lower() for a in alternatives)}
        for term in group:
            index.setdefault(term, set()).update(group)
    return index


SYNONYM_INDEX = _build_synonym_index()


def _wordnet_synonyms(term):
    """WordNet synonyms for a term, when NLTK and its corpus are both present.

    Layered *under* the domain map above, never over it: WordNet is useful for
    ordinary verbs ("buy"/"purchase") but knows nothing about Kubernetes, and
    would happily offer "conduce" as a synonym for "lead". Multi-word technical
    terms are left alone.
    """
    if _wordnet is None or ' ' in term or len(term) < 4:
        return set()
    try:
        forms = set()
        for synset in _wordnet.synsets(term):
            # Only verbs and nouns; adjective synsets are too loose to be useful.
            if synset.pos() not in ('v', 'n'):
                continue
            for lemma in synset.lemmas():
                name = lemma.name().replace('_', ' ').lower()
                if name != term.lower() and ' ' not in name:
                    forms.add(name)
        return forms
    except Exception:  # pragma: no cover - corpus lookup issues
        return set()


def _synonyms_for(term):
    """All accepted surface forms of ``term`` (including itself)."""
    key = term.lower()
    if key in SYNONYM_INDEX:
        # A curated entry is authoritative — don't dilute it with WordNet.
        return SYNONYM_INDEX[key]
    return {key} | _wordnet_synonyms(key)


def _contains_term(haystack_lower, term):
    """Word-boundary containment test that tolerates regex-special skill names."""
    pattern = r'(?<!\w)' + re.escape(term.lower()) + r'(?!\w)'
    return re.search(pattern, haystack_lower) is not None


def _count_term(haystack_lower, term):
    pattern = r'(?<!\w)' + re.escape(term.lower()) + r'(?!\w)'
    return len(re.findall(pattern, haystack_lower))


def _tokens(text):
    return re.findall(r"[a-z0-9+#./]+", (text or '').lower())


def _sentences(text):
    """Split into bullet points and sentences — the units an ATS checks proximity in."""
    units = []
    for line in (text or '').splitlines():
        line = line.strip()
        if not line:
            continue
        if line.lstrip().startswith(BULLET_PREFIXES):
            units.append(line)
        else:
            units.extend(s.strip() for s in re.split(r'(?<=[.!?])\s+', line) if s.strip())
    return units


def _expected_frequency(jd_count):
    """How often a CV should mention a term the job description uses ``jd_count`` times.

    An ATS reads a term the job leans on repeatedly as a core requirement: if the
    advert says "SQL" five times, one passing mention on the CV reads as a nice-
    to-have. Roughly half the job's usage, capped at 3 (beyond that it is
    stuffing, not emphasis).
    """
    if jd_count <= 1:
        return 1
    return min(3, max(1, round(jd_count / 2)))


def _within_window(tokens, words, window=PROXIMITY_WINDOW):
    """Do all of ``words`` appear within a ``window``-token span of ``tokens``?

    This is the real proximity test: "Financial" and "Forecasting" scattered
    across different sections score nothing, but "financial forecasting for 12
    regions" scores, because the terms sit inside one bullet, near each other.
    """
    stems = [_stem(w) for w in words]
    positions = []
    for stem in stems:
        hits = [i for i, t in enumerate(tokens) if _stem(t) == stem]
        if not hits:
            return False
        positions.append(hits)

    # Slide over every combination cheaply: for each occurrence of the first term,
    # check that every other term has an occurrence inside the window around it.
    for anchor in positions[0]:
        low, high = anchor - window, anchor + window
        if all(any(low <= h <= high for h in hits) for hits in positions[1:]):
            return True
    return False


def _fuzzy_ratio(a, b):
    """0-100 similarity. stdlib difflib — same algorithm fuzzywuzzy wraps."""
    from difflib import SequenceMatcher

    if not a or not b:
        return 0
    return int(round(SequenceMatcher(None, a.lower(), b.lower()).ratio() * 100))


def _clamp(value, low=0, high=100):
    return max(low, min(high, int(round(value))))


# ---------------------------------------------------------------------------
# Date / experience parsing
# ---------------------------------------------------------------------------

def _parse_date_token(token):
    """Parse "Jan 2022" / "01/2022" / "2022" / "Present" into a (year, month) tuple."""
    token = token.strip().lower().rstrip('.')
    if token in ('present', 'current', 'now', 'ongoing'):
        today = date.today()
        return (today.year, today.month)

    month_match = re.match(r'([a-z]{3,9})\.?\s+(\d{4})', token)
    if month_match:
        month = MONTHS.get(month_match.group(1)[:3])
        if month:
            return (int(month_match.group(2)), month)
        return (int(month_match.group(2)), 1)

    slash_match = re.match(r'(\d{1,2})/(\d{4})', token)
    if slash_match:
        month = _clamp(int(slash_match.group(1)), 1, 12)
        return (int(slash_match.group(2)), month)

    year_match = re.match(r'(\d{4})', token)
    if year_match:
        return (int(year_match.group(1)), 1)
    return None


def _months_between(start, end):
    return max(0, (end[0] - start[0]) * 12 + (end[1] - start[1]))


def extract_date_ranges(text):
    """Return [(start, end)] as (year, month) tuples, in the order they appear."""
    ranges = []
    for match in DATE_RANGE_RE.finditer(text or ''):
        start = _parse_date_token(match.group('start'))
        end = _parse_date_token(match.group('end'))
        if not (start and end):
            continue
        # A reversed range is a parsing artefact, not real history — skip it.
        if _months_between(start, end) <= 0 and end != start:
            continue
        ranges.append((start, end))
    return ranges


def total_experience_years(text):
    """Total years of experience, merging overlapping roles so they aren't double-counted."""
    ranges = extract_date_ranges(text)
    if not ranges:
        return 0.0

    merged = []
    for start, end in sorted(ranges):
        if merged and start <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))

    months = sum(_months_between(s, e) for s, e in merged)
    return round(months / 12.0, 1)


# ---------------------------------------------------------------------------
# Job-description requirement extraction (drives the Phase 2 knock-outs)
# ---------------------------------------------------------------------------

def extract_job_requirements(job_description, job_title='', job_location=''):
    """Mine the hard requirements an ATS would knock candidates out on.

    Every key is optional: a requirement the job description never states cannot
    knock anybody out, so it is left as None/empty and the corresponding check
    is skipped.
    """
    text = job_description or ''
    lowered = text.lower()

    years = [int(m.group(1)) for m in YEARS_REQUIRED_RE.finditer(text)]
    # Several figures usually means "3 years in X, 5 in Y" — the highest is the bar.
    required_years = max(years) if years else None

    required_education = None
    for level, label, keywords in DEGREE_LEVELS:
        if level < 3:
            continue  # only degree-level requirements are true knock-outs
        if any(_contains_term(lowered, kw) for kw in keywords):
            # Take the *lowest* degree mentioned: "BSc or MSc" means a BSc suffices.
            if required_education is None or level < required_education[0]:
                required_education = (level, label)
    required_education = required_education[1] if required_education else None

    required_certs = [c for c in CERTIFICATIONS if _contains_term(lowered, c)]

    required_gpa = None
    gpa_match = GPA_REQUIRED_RE.search(text)
    if gpa_match:
        required_gpa = float(gpa_match.group(1) or gpa_match.group(2))
    else:
        # UK adverts state a classification ("2:1 or above"), not a GPA.
        for phrase, value in UK_CLASSIFICATIONS.items():
            if re.search(
                r'(?:minimum|min\.?|at least|or above|or higher)[^.]{0,20}'
                + re.escape(phrase) + r'|' + re.escape(phrase)
                + r'[^.]{0,20}(?:or above|or higher|minimum)',
                lowered,
            ):
                required_gpa = value
                break

    requires_auth = None
    for label, pattern in WORK_AUTH_PATTERNS:
        if re.search(pattern, lowered):
            requires_auth = label
            break

    return {
        'title': job_title or '',
        'required_years': required_years,
        'required_education': required_education,
        'required_certifications': required_certs,
        'required_gpa': required_gpa,
        'required_location': job_location or '',
        'work_authorization': requires_auth,
    }


# ---------------------------------------------------------------------------
# Keyword extraction from the job description
# ---------------------------------------------------------------------------
# Weight by keyword class: an ATS ranks a missing hard skill far above a missing
# adjective.
WEIGHT_HARD = 3.0
WEIGHT_CERT = 3.0
WEIGHT_GENERAL = 2.0
WEIGHT_SOFT = 1.0


def extract_jd_keywords(job_description, limit=60):
    """Return the ranked keywords an ATS would screen this CV against.

    Each item: {'term', 'type', 'weight', 'jd_count'}. Hard skills and required
    certifications outrank generic phrases, which outrank soft skills.
    """
    text = job_description or ''
    lowered = text.lower()
    keywords = []
    seen = set()

    def add(term, kind, weight):
        key = term.lower()
        if key in seen or len(key) < 2:
            return
        seen.add(key)
        keywords.append({
            'term': term,
            'type': kind,
            'weight': weight,
            'jd_count': _count_term(lowered, term),
        })

    for skill in HARD_SKILL_VOCAB:
        if _contains_term(lowered, skill):
            add(skill, 'hard', WEIGHT_HARD)
    for cert in CERTIFICATIONS:
        if _contains_term(lowered, cert):
            add(cert, 'certification', WEIGHT_CERT)
    for soft in SOFT_SKILLS:
        if _contains_term(lowered, soft):
            add(soft, 'soft', WEIGHT_SOFT)

    # Remaining signal: the most frequent meaningful uni-/bi-grams, which is how
    # an ATS picks up domain terms that aren't in any predefined vocabulary.
    tokens = _tokens(text)
    counts = {}
    for size in (1, 2):
        for i in range(len(tokens) - size + 1):
            gram = tokens[i:i + size]
            if any(t in STOPWORDS or len(t) < 3 or t.isdigit() for t in gram):
                continue
            phrase = ' '.join(gram)
            counts[phrase] = counts.get(phrase, 0) + 1

    ranked = sorted(counts.items(), key=lambda kv: (-kv[1], kv[0]))
    for phrase, count in ranked:
        if len(keywords) >= limit:
            break
        if count < 2:  # a term mentioned once isn't a screening keyword
            break
        add(phrase, 'general', WEIGHT_GENERAL)

    return keywords[:limit]


# ---------------------------------------------------------------------------
# File-level inspection (Phase 1)
# ---------------------------------------------------------------------------

def _load_pdfplumber():
    try:
        import pdfplumber

        return pdfplumber
    except ImportError:  # pragma: no cover - optional dependency
        logger.info('pdfplumber not installed; PDF layout checks will be skipped.')
        return None


def _inspect_docx(path):
    """Layout, prohibited elements, fonts and header/footer text from a DOCX."""
    import docx

    document = docx.Document(path)
    findings = {
        'prohibited': [],
        'fonts': set(),
        'header_footer_text': '',
        'columns': 1,
        'available': True,
    }

    if document.tables:
        findings['prohibited'].append(f'{len(document.tables)} table(s)')
    if document.inline_shapes:
        findings['prohibited'].append(f'{len(document.inline_shapes)} image(s)/shape(s)')

    body_xml = document.element.body.xml
    if 'txbxContent' in body_xml:
        findings['prohibited'].append('text box(es)')
    if '<w:drawing' in body_xml or '<w:pict' in body_xml:
        if not any('image' in p for p in findings['prohibited']):
            findings['prohibited'].append('embedded drawing(s)/image(s)')

    # Multi-column sections are declared in the section properties.
    for match in re.finditer(r'w:num="(\d+)"', body_xml):
        findings['columns'] = max(findings['columns'], int(match.group(1)))

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.font.name:
                findings['fonts'].add(run.font.name)
    try:
        style_font = document.styles['Normal'].font.name
        if style_font:
            findings['fonts'].add(style_font)
    except Exception:
        pass

    header_footer = []
    for section in document.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                if paragraph.text.strip():
                    header_footer.append(paragraph.text.strip())
    findings['header_footer_text'] = '\n'.join(header_footer)
    return findings


def _inspect_pdf(path):
    """Layout, prohibited elements and fonts from a PDF, via pdfplumber."""
    pdfplumber = _load_pdfplumber()
    if pdfplumber is None:
        return {'available': False}

    findings = {
        'prohibited': [],
        'fonts': set(),
        'header_footer_text': '',
        'columns': 1,
        'available': True,
    }
    image_count = 0
    table_count = 0

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            image_count += len(page.images or [])
            try:
                table_count += len(page.find_tables() or [])
            except Exception:
                pass
            for char in (page.chars or []):
                name = (char.get('fontname') or '')
                # Subset prefixes look like "ABCDEF+Arial-Bold".
                name = re.sub(r'^[A-Z]{6}\+', '', name)
                name = re.sub(r'[-,](?:Bold|Italic|Oblique|Regular|Light|Medium|BoldItalic|MT|PS).*$', '', name)
                if name:
                    findings['fonts'].add(name)
            if _page_is_multicolumn(page):
                findings['columns'] = max(findings['columns'], 2)

    if image_count:
        findings['prohibited'].append(f'{image_count} image(s)')
    if table_count:
        findings['prohibited'].append(f'{table_count} table(s)')
    return findings


def _page_is_multicolumn(page):
    """Detect a two-column layout from a vertical gutter in the word x-positions.

    A single-column CV has words spanning the full width; a two-column one has a
    band in the middle that no word crosses, with substantial text either side.
    """
    try:
        words = page.extract_words() or []
    except Exception:
        return False
    if len(words) < 40:
        return False

    width = page.width or 1
    # Scan candidate gutters across the middle half of the page.
    for split in [width * f for f in (0.35, 0.4, 0.45, 0.5, 0.55, 0.6, 0.65)]:
        crossing = sum(1 for w in words if w['x0'] < split < w['x1'])
        left = sum(1 for w in words if w['x1'] <= split)
        right = sum(1 for w in words if w['x0'] >= split)
        if crossing == 0 and left >= len(words) * 0.2 and right >= len(words) * 0.2:
            return True
    return False


# ---------------------------------------------------------------------------
# Heading detection
# ---------------------------------------------------------------------------

def _normalize_heading(line):
    cleaned = line.strip().strip('#*_ ').rstrip(':').strip()
    cleaned = re.sub(r'\s+', ' ', cleaned).replace('&', 'and')
    return cleaned.lower()


def find_headings(cv_text):
    """Split the CV's heading-looking lines into recognised and non-standard ones.

    A heading is a short line with no terminal punctuation. Recognised ones map
    onto the canonical sections an ATS knows; anything else ("My Journey") is a
    creative title an ATS would fail to categorise.
    """
    standard, creative = [], []
    lines = [l.strip() for l in (cv_text or '').splitlines()]
    non_empty = [l for l in lines if l]
    # The first line is the candidate's name, never a section heading.
    name_line = non_empty[0] if non_empty else ''

    for line in non_empty:
        if len(line.split()) > 6 or line == name_line:
            continue
        if line.lstrip().startswith(BULLET_PREFIXES):
            continue

        normalized = _normalize_heading(line)
        if normalized in SECTION_ALIASES:
            standard.append(line.strip('#*_ ').strip())
            continue

        # A creative heading is a short line that *reads* like a heading: no
        # sentence-ending punctuation, and starting with a capital. Contact,
        # date and "Title | Company" lines are excluded.
        if re.search(r'[@|]|\d{4}|\+\d|www\.|\.com', line):
            continue
        if line.rstrip(':').endswith(('.', ',', ';', '!', '?')):
            continue
        stripped = line.strip('#*_ ').rstrip(':').strip()
        if len(stripped) < 3 or not stripped[0].isupper():
            continue
        creative.append(stripped)
    return standard, creative


# ---------------------------------------------------------------------------
# The checker
# ---------------------------------------------------------------------------

class ATSChecker:
    """Score a CV against a job description the way an ATS would.

    ``cv_text``          plain text of the CV (already extracted).
    ``job_description``  the job advert text.
    ``job_requirements`` optional dict of hard requirements; when omitted they
                         are mined from the job description itself.
    ``file_path``        optional path to the original CV file. Without it the
                         Phase 1 file checks are skipped (headers are still
                         checked, as those come from the text).
    """

    def __init__(self, cv_text, job_description, job_requirements=None, file_path=None):
        self.cv_text = cv_text or ''
        self.job_description = job_description or ''
        self.file_path = file_path
        self.job_requirements = (
            job_requirements
            if job_requirements is not None
            else extract_job_requirements(self.job_description)
        )

        self.cv_lower = self.cv_text.lower()
        self.sections = parse_cv_sections(self.cv_text)
        self.keywords = extract_jd_keywords(self.job_description)
        self.results = {}
        self._recommendations = []

    # -- helpers ------------------------------------------------------------

    def _section_text(self, key):
        value = self.sections.get(key)
        if isinstance(value, list):
            return '\n'.join(value)
        return value or ''

    @property
    def experience_text(self):
        return self._section_text('experience')

    @property
    def work_history_text(self):
        """The text employment dates should be read from.

        Deliberately excludes the education section: "BSc 2015 - 2018" is three
        years of study, not three years of work, and counting it would let an
        underqualified CV pass the years knock-out.
        """
        experience = self.experience_text
        if experience:
            return experience

        # No parsable experience section — fall back to the whole CV, minus any
        # education lines we can identify.
        education_lines = set(self.sections.get('education') or [])
        if not education_lines:
            return self.cv_text
        return '\n'.join(
            line for line in self.cv_text.splitlines()
            if line.strip() not in education_lines
        )

    @property
    def recent_experience_text(self):
        """Text of the two most recent roles — what recency weighting looks at."""
        lines = self.sections.get('experience') or []
        # Roles start at a "Title | Company | Location" line; take the first two blocks.
        blocks, current = [], []
        for line in lines:
            if '|' in line and current:
                blocks.append(current)
                current = [line]
            else:
                current.append(line)
        if current:
            blocks.append(current)
        return '\n'.join('\n'.join(b) for b in blocks[:2]) or '\n'.join(lines[:15])

    def _found_in(self, text_lower, term):
        """Does ``term`` (or a synonym, or its stem) appear in the given text?"""
        for form in _synonyms_for(term):
            if _contains_term(text_lower, form):
                return True
        # Stem fallback, so "managing projects" satisfies "project management".
        stem_key = _stem_phrase(term)
        if not stem_key:
            return False
        text_stems = ' '.join(_stem(t) for t in _tokens(text_lower))
        return stem_key in text_stems

    def _recommend(self, message):
        if message not in self._recommendations:
            self._recommendations.append(message)

    # -- Phase 1: parsing & technical formatting ----------------------------

    def check_file_format(self, file_path=None):
        """Is the file machine-readable? Image-based PDFs yield almost no text."""
        path = file_path or self.file_path
        if not path:
            return {'pass': True, 'format': 'TEXT', 'skipped': True}

        ext = os.path.splitext(path)[1].lower()
        result = {'format': ext.lstrip('.').upper() or 'UNKNOWN', 'skipped': False}

        if ext not in ('.docx', '.pdf'):
            result['pass'] = False
            result['issue'] = f'Unsupported file type "{ext}". Use DOCX or PDF.'
            self._recommend('Save your CV as a DOCX or a text-based PDF.')
            return result

        # Under ~100 characters the file is a scan/image, and an ATS reads nothing.
        if len(self.cv_text.strip()) < 100:
            result['pass'] = False
            result['issue'] = (
                'Almost no text could be extracted — the CV is likely image-based '
                'or scanned. An ATS would read it as empty.'
            )
            self._recommend(
                'Your CV appears to be an image or scan. Rebuild it as a real text '
                'document — an ATS cannot read text inside an image.'
            )
            return result

        result['pass'] = True
        if ext == '.pdf':
            result['note'] = 'PDF is accepted, but DOCX parses most reliably.'
        return result

    def check_layout(self, file_path=None):
        """Single-column, top-to-bottom? Multi-column layouts scramble the parse."""
        findings = self._file_findings(file_path)
        if not findings.get('available'):
            return {'pass': True, 'layout': 'unknown', 'skipped': True}

        columns = findings.get('columns', 1)
        if columns > 1:
            self._recommend(
                'Reformat your CV into a single column. Multi-column layouts are '
                'read left-to-right across both columns and come out scrambled.'
            )
            return {'pass': False, 'layout': f'{columns}-column', 'skipped': False}
        return {'pass': True, 'layout': 'single-column', 'skipped': False}

    def check_prohibited_elements(self, file_path=None):
        """Tables, images, text boxes — content an ATS silently drops."""
        findings = self._file_findings(file_path)
        if not findings.get('available'):
            return {'pass': True, 'elements': [], 'skipped': True}

        elements = findings.get('prohibited', [])
        if elements:
            self._recommend(
                'Remove ' + ', '.join(elements) + '. Content inside tables, text '
                'boxes and images is usually dropped entirely by an ATS parser.'
            )
        return {'pass': not elements, 'elements': elements, 'skipped': False}

    def check_header_footer(self, file_path=None):
        """Contact details in a header/footer are invisible to most parsers."""
        findings = self._file_findings(file_path)
        if not findings.get('available'):
            return {'pass': True, 'skipped': True}

        header_footer = (findings.get('header_footer_text') or '').strip()
        if not header_footer:
            return {'pass': True, 'skipped': False, 'contains_contact': False}

        # Only a problem if the contact details appear *only* there.
        has_contact = bool(re.search(r'[\w.+-]+@[\w-]+\.\w+|\+?\d[\d\s()-]{7,}', header_footer))
        body_has_contact = bool(
            re.search(r'[\w.+-]+@[\w-]+\.\w+|\+?\d[\d\s()-]{7,}', self.cv_text)
        )
        if has_contact and not body_has_contact:
            self._recommend(
                'Move your name, phone number and email out of the page header/footer '
                'and into the body of the CV — parsers ignore headers and footers.'
            )
            return {'pass': False, 'skipped': False, 'contains_contact': True}
        return {'pass': True, 'skipped': False, 'contains_contact': has_contact}

    def check_fonts(self, file_path=None):
        """Non-standard fonts can garble the extracted text."""
        findings = self._file_findings(file_path)
        if not findings.get('available'):
            return {'pass': True, 'fonts': [], 'issues': [], 'skipped': True}

        fonts = sorted(findings.get('fonts') or [])
        issues = [f for f in fonts if f.lower().strip() not in STANDARD_FONTS]
        if issues:
            self._recommend(
                'Switch to a standard font (Arial, Calibri, Times New Roman or '
                'Verdana). Non-standard fonts can be extracted as garbled text: '
                + ', '.join(issues[:5])
            )

        garbled = self._garbled_ratio()
        if garbled > 0.01:
            issues.append(f'text extracts as garbled characters ({garbled:.0%})')
            self._recommend(
                'Your CV\'s text extracts as garbled symbols, which means the font '
                'is not being decoded. Re-save it in a standard font — an ATS would '
                'read gibberish.'
            )
        return {
            'pass': not issues, 'fonts': fonts, 'issues': issues,
            'garbled_ratio': round(garbled, 4), 'skipped': False,
        }

    def _garbled_ratio(self):
        """Fraction of extracted characters that came out as replacement/control junk.

        A custom font that the parser cannot decode yields U+FFFD ("�") or
        control characters — the ATS reads gibberish, whatever the CV looks like
        on screen.
        """
        if not self.cv_text:
            return 0.0
        bad = sum(
            1 for ch in self.cv_text
            if ch == '�' or (ord(ch) < 32 and ch not in '\n\r\t')
        )
        return bad / len(self.cv_text)

    def check_section_headers(self):
        """Standard headings only — an ATS files content by the headings it knows."""
        standard, creative = find_headings(self.cv_text)
        found_keys = {
            SECTION_ALIASES[_normalize_heading(h)]
            for h in standard
            if _normalize_heading(h) in SECTION_ALIASES
        }
        missing = sorted(CRITICAL_SECTIONS - found_keys)

        if missing:
            self._recommend(
                'Add the missing section heading(s): '
                + ', '.join(m.title() for m in missing)
                + '. An ATS files your content by heading, and discards what it '
                'cannot categorise.'
            )
        if creative:
            self._recommend(
                'Rename the non-standard heading(s) '
                + ', '.join(f'"{c}"' for c in creative[:4])
                + ' to standard ones (Summary, Professional Experience, Education, '
                'Skills, Certifications).'
            )

        return {
            'pass': not missing,
            'standard_headers': standard,
            'missing_headers': [m.title() for m in missing],
            'creative_headers': creative,
        }

    def _file_findings(self, file_path=None):
        """Inspect the CV file once and cache the result."""
        path = file_path or self.file_path
        if not path or not os.path.exists(path):
            return {'available': False}
        if getattr(self, '_cached_findings', None) is not None:
            return self._cached_findings

        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == '.docx':
                findings = _inspect_docx(path)
            elif ext == '.pdf':
                findings = _inspect_pdf(path)
            else:
                findings = {'available': False}
        except Exception:
            logger.exception('ATS file inspection failed for %s', path)
            findings = {'available': False}

        self._cached_findings = findings
        return findings

    def run_phase1(self):
        """Parsing & technical formatting. A failure here is disqualifying."""
        file_format = self.check_file_format()
        layout = self.check_layout()
        prohibited = self.check_prohibited_elements()
        header_footer = self.check_header_footer()
        fonts = self.check_fonts()
        headers = self.check_section_headers()

        checks = [file_format, layout, prohibited, header_footer, fonts, headers]
        passed = sum(1 for c in checks if c.get('pass'))
        score = _clamp(passed / len(checks) * 100)

        result = {
            # File format and section headers are the disqualifying ones; a
            # non-standard font alone shouldn't reject an otherwise readable CV.
            'pass': file_format.get('pass', True) and headers.get('pass', True)
                    and layout.get('pass', True) and prohibited.get('pass', True),
            'score': score,
            'file_format': file_format.get('format'),
            'file_format_ok': file_format.get('pass'),
            'layout': layout.get('layout'),
            'prohibited_elements': prohibited.get('elements', []),
            'standard_headers': headers['standard_headers'],
            'missing_headers': headers['missing_headers'],
            'creative_headers': headers['creative_headers'],
            'fonts': fonts.get('fonts', []),
            'font_issues': fonts.get('issues', []),
            'header_footer_ok': header_footer.get('pass'),
            'file_checks_skipped': file_format.get('skipped') or layout.get('skipped'),
        }
        if file_format.get('issue'):
            result['issue'] = file_format['issue']
        self.results['phase1_parsing'] = result
        return result

    # -- Phase 2: knock-out filters -----------------------------------------

    def check_experience_years(self, required_years=None):
        required = required_years if required_years is not None \
            else self.job_requirements.get('required_years')
        found = total_experience_years(self.work_history_text)
        if required is None:
            return {'required': None, 'found': found, 'pass': True, 'skipped': True}

        if found == 0:
            # No parsable dates at all. That is a CV problem, not proof the
            # candidate is unqualified — flag it, but never knock them out on it.
            self._recommend(
                'No employment dates could be read from your CV. Add them as '
                '"Jan 2020 - Mar 2023" under each role, or an ATS cannot verify '
                'your years of experience.'
            )
            return {
                'required': required, 'found': 0.0, 'pass': True,
                'skipped': True, 'unverifiable': True,
            }

        # Half a year of grace: date parsing rounds, and so do recruiters.
        passed = found >= (required - 0.5)
        if not passed:
            self._recommend(
                f'This role asks for {required} years of experience; your CV shows '
                f'{found}. Make sure every role\'s dates are present and formatted '
                f'as "Jan 2020 - Mar 2023".'
            )
        return {'required': required, 'found': found, 'pass': passed, 'skipped': False}

    def check_education_level(self, required_level=None):
        required = required_level if required_level is not None \
            else self.job_requirements.get('required_education')
        found_level, found_label = self._highest_degree()
        if not required:
            return {
                'required': None, 'found': found_label, 'pass': True, 'skipped': True,
            }

        required_rank = next(
            (lvl for lvl, label, _ in DEGREE_LEVELS if label == required), 3
        )
        passed = found_level >= required_rank
        if not passed:
            self._recommend(
                f'This role requires a {required} degree; the CV shows '
                f'{found_label or "no degree"}. If you hold one, state it explicitly '
                f'under an EDUCATION heading.'
            )
        return {
            'required': required, 'found': found_label, 'pass': passed, 'skipped': False,
        }

    def _highest_degree(self):
        """Highest degree on the CV as (rank, label)."""
        education = self._section_text('education').lower() or self.cv_lower
        for level, label, keywords in DEGREE_LEVELS:
            if any(_contains_term(education, kw) for kw in keywords):
                return level, label
        return 0, ''

    def check_certifications(self, required_certs=None):
        required = required_certs if required_certs is not None \
            else self.job_requirements.get('required_certifications') or []
        if not required:
            return {'required': [], 'found': [], 'missing': [], 'pass': True, 'skipped': True}

        found = [c for c in required if _contains_term(self.cv_lower, c)]
        missing = [c for c in required if c not in found]
        if missing:
            self._recommend(
                'The job requires these certification(s): '
                + ', '.join(c.upper() for c in missing)
                + '. If you hold them, list them under a CERTIFICATIONS heading — '
                'an ATS looks for the exact name.'
            )
        return {
            'required': required, 'found': found, 'missing': missing,
            'pass': not missing, 'skipped': False,
        }

    def check_location(self, required_location=None):
        required = required_location if required_location is not None \
            else self.job_requirements.get('required_location') or ''
        if not required:
            return {'required': '', 'found': '', 'pass': True, 'skipped': True}

        # Compare on the place names, not the whole "London, England, UK" string.
        parts = [p.strip().lower() for p in re.split(r'[,/]', required) if p.strip()]
        matched = any(_contains_term(self.cv_lower, p) for p in parts if len(p) > 2)
        remote = _contains_term(required.lower(), 'remote') or _contains_term(
            self.job_description.lower(), 'fully remote'
        )
        willing = bool(RELOCATION_RE.search(self.cv_text))
        cv_location = self._cv_location()

        if matched or remote or willing:
            return {
                'required': required, 'found': cv_location or required,
                'relocation_stated': willing, 'pass': True, 'skipped': False,
            }

        if not cv_location:
            # We could not read a location off the CV. That is not evidence the
            # candidate is in the wrong place, so it must not knock them out —
            # otherwise every CV without a parsable address is rejected.
            self._recommend(
                f'The role is based in {required}, but no location could be read '
                f'from your CV. Add "City, Country" to your contact line (and '
                f'"Willing to relocate" if that applies).'
            )
            return {
                'required': required, 'found': '', 'relocation_stated': False,
                'pass': True, 'skipped': True, 'unverifiable': True,
            }

        # A location was found and it genuinely conflicts — a real knock-out.
        self._recommend(
            f'The role is based in {required} but your CV shows {cv_location}. '
            f'Add "Willing to relocate" if you are.'
        )
        return {
            'required': required, 'found': cv_location, 'relocation_stated': False,
            'pass': False, 'skipped': False,
        }

    def _cv_location(self):
        """Best guess at the candidate's location: the contact line, usually."""
        contact = self.sections.get('contact') or ''
        header = f"{self.sections.get('name', '')} {contact}"
        match = re.search(
            r'([A-Z][a-zA-Z]+(?:\s[A-Z][a-zA-Z]+)?,\s*(?:UK|United Kingdom|England|'
            r'Scotland|Wales|USA|US|United States))',
            header,
        )
        return match.group(1) if match else ''

    def _cv_gpa(self):
        """The CV's GPA, or a UK classification (2:1, First) mapped onto a 4.0 scale."""
        education_text = self._section_text('education')
        gpa_match = GPA_RE.search(education_text or self.cv_text)
        if gpa_match:
            return float(gpa_match.group(1))
        lowered = (education_text or self.cv_text).lower()
        for phrase, value in UK_CLASSIFICATIONS.items():
            if _contains_term(lowered, phrase):
                return value
        return None

    def check_gpa(self, required_gpa=None):
        """Phase 2: does the CV meet the job's minimum GPA / degree classification?"""
        required = required_gpa if required_gpa is not None \
            else self.job_requirements.get('required_gpa')
        found = self._cv_gpa()
        if required is None:
            return {'required': None, 'found': found, 'pass': True, 'skipped': True}

        if found is None:
            # The job wants a GPA and the CV states none. That is a CV omission,
            # not proof of a low grade, so it must not knock the candidate out.
            self._recommend(
                f'This role asks for a minimum grade ({required} GPA or equivalent) '
                f'and your CV states none. Add your GPA or degree classification '
                f'(e.g. "2:1") to the education section.'
            )
            return {
                'required': required, 'found': None, 'pass': True,
                'skipped': True, 'unverifiable': True,
            }

        passed = found >= required
        if not passed:
            self._recommend(
                f'This role requires a minimum GPA of {required}; your CV shows '
                f'{found}. This is a hard filter for this job.'
            )
        return {
            'required': required, 'found': found, 'pass': passed, 'skipped': False,
        }

    def check_work_authorization(self):
        required = self.job_requirements.get('work_authorization')
        found = None
        for label, pattern in WORK_AUTH_PATTERNS:
            if re.search(pattern, self.cv_lower):
                found = label
                break

        if not required:
            return {'required': None, 'found': found, 'pass': True, 'skipped': True}

        # The only genuine knock-out: the job won't sponsor and the CV needs it.
        needs_sponsorship = found == 'Requires Sponsorship'
        job_wont_sponsor = required in ('Citizen', 'Permanent Resident', 'Authorised to Work')
        passed = not (needs_sponsorship and job_wont_sponsor)
        if not passed:
            self._recommend(
                'This employer does not sponsor visas and your CV states that you '
                'require sponsorship. This is a hard knock-out for this role.'
            )
        return {'required': required, 'found': found, 'pass': passed, 'skipped': False}

    def run_phase2(self):
        """Knock-out filters. Any failure here is disqualifying."""
        experience = self.check_experience_years()
        education = self.check_education_level()
        certifications = self.check_certifications()
        location = self.check_location()
        authorization = self.check_work_authorization()
        gpa = self.check_gpa()

        checks = {
            'experience_years': experience,
            'education': education,
            'certifications': certifications,
            'gpa': gpa,
            'location': location,
            'work_authorization': authorization,
        }
        failed = [name for name, c in checks.items() if not c['pass']]
        result = {'pass': not failed, 'failed_filters': failed, **checks}
        self.results['phase2_knockout'] = result
        return result

    # -- Phase 3: keyword matching ------------------------------------------

    def calculate_keyword_match(self):
        """Weighted keyword coverage, with a bonus for prominent placement."""
        if not self.keywords:
            result = {
                'score': 50, 'exact_matches': 0, 'total_keywords': 0,
                'match_percentage': 0, 'hard_skills_found': [],
                'hard_skills_missing': [], 'soft_skills_found': [],
                'missing_keywords': [], 'keyword_density': {},
                'keyword_stuffing': False,
                'note': 'No keywords could be mined from the job description.',
            }
            self.results['phase3_keyword'] = result
            return result

        summary_text = self._section_text('profile').lower()
        recent_text = self.recent_experience_text.lower()

        earned = 0.0
        possible = 0.0
        exact_matches = 0
        matched, missing = [], []
        hard_found, hard_missing, soft_found = [], [], []
        density = {}
        stuffed = []
        underused = []

        for keyword in self.keywords:
            term = keyword['term']
            weight = keyword['weight']
            possible += weight

            found = self._found_in(self.cv_lower, term)
            count = _count_term(self.cv_lower, term)
            if count:
                exact_matches += 1

            if found:
                # Contextual placement (the TF-IDF-ish part): a keyword in the
                # summary or the most recent role counts for more than one buried
                # in a 2011 role.
                multiplier = 1.0
                if self._found_in(summary_text, term):
                    multiplier = 1.25
                elif self._found_in(recent_text, term):
                    multiplier = 1.15

                # Frequency sufficiency: a term the job leans on repeatedly needs
                # to appear more than once here to read as a core skill rather
                # than a passing mention.
                expected = _expected_frequency(keyword['jd_count'])
                if count and count < expected:
                    multiplier *= 0.75
                    underused.append({
                        'term': term, 'cv_count': count, 'expected': expected,
                        'jd_count': keyword['jd_count'],
                    })

                earned += weight * multiplier
                matched.append(term)
                density[term] = count

                if keyword['type'] == 'hard':
                    hard_found.append(term)
                elif keyword['type'] == 'soft':
                    soft_found.append(term)

                # Stuffing: far more repetitions than the job description itself.
                if count > max(5, keyword['jd_count'] * 3):
                    stuffed.append(term)
            else:
                missing.append(term)
                if keyword['type'] in ('hard', 'certification'):
                    hard_missing.append(term)

        score = _clamp(earned / possible * 100) if possible else 0
        match_percentage = _clamp(len(matched) / len(self.keywords) * 100)

        if stuffed:
            score = _clamp(score - 10)
            self._recommend(
                'Reduce repetition of ' + ', '.join(stuffed[:3])
                + '. Modern ATS platforms penalise keyword stuffing.'
            )
        if hard_missing:
            self._recommend(
                'Add these skills from the job description, if you genuinely have '
                'them: ' + ', '.join(hard_missing[:6]) + '.'
            )
        if underused:
            worst = sorted(underused, key=lambda u: u['jd_count'], reverse=True)[:3]
            self._recommend(
                'These terms are central to the job but barely appear on your CV: '
                + ', '.join(
                    f'"{u["term"]}" (job mentions it {u["jd_count"]}x, your CV '
                    f'{u["cv_count"]}x — aim for {u["expected"]})' for u in worst
                )
                + '. Work them into the roles where you genuinely used them.'
            )

        result = {
            'score': score,
            'exact_matches': exact_matches,
            'total_keywords': len(self.keywords),
            'matched_keywords': matched,
            'match_percentage': match_percentage,
            'hard_skills_found': hard_found,
            'hard_skills_missing': hard_missing,
            'soft_skills_found': soft_found,
            'missing_keywords': missing,
            'keyword_density': density,
            'keyword_stuffing': bool(stuffed),
            'stuffed_keywords': stuffed,
            'underused_keywords': underused,
        }
        self.results['phase3_keyword'] = result
        return result

    # -- Phase 4: context & proximity ---------------------------------------

    def check_job_title_match(self):
        """Does the CV carry job titles resembling the target title?"""
        target = (self.job_requirements.get('title') or '').strip()
        if not target:
            return 50, []

        titles = []
        for line in (self.sections.get('experience') or []):
            if '|' in line:
                titles.append(line.split('|')[0].strip())
        if not titles:
            # No structured roles: fall back to the CV's first line (often a title).
            titles = [self.sections.get('name', '')]

        best = max((_fuzzy_ratio(target, t) for t in titles if t), default=0)
        if best < 60:
            self._recommend(
                f'None of your job titles closely match "{target}". Where it is '
                f'honest to do so, align a title (e.g. add the target title in '
                f'brackets after your official one).'
            )
        return best, titles

    def check_proximity(self):
        """Are the job's skills evidenced inside experience bullets, or just listed?

        Two things an ATS checks, and this is what separates a real practitioner
        from someone who pasted a skills list:

        1. Evidence — does the skill appear in the experience section at all, or
           only in the skills list?
        2. Proximity — for a multi-word requirement like "financial forecasting",
           do those words actually sit together inside one bullet point, rather
           than being scattered across the CV?
        """
        phase3 = self.results.get('phase3_keyword') or self.calculate_keyword_match()
        hard_found = phase3.get('hard_skills_found', [])
        if not hard_found:
            return 50

        experience_lower = self.experience_text.lower()
        if not experience_lower:
            self._recommend(
                'Your skills are listed but never evidenced. Work the job\'s key '
                'skills into your experience bullet points, where an ATS checks '
                'that they co-occur with real achievements.'
            )
            return 0

        # 1. Evidence.
        evidenced = [s for s in hard_found if self._found_in(experience_lower, s)]
        evidence_score = _clamp(len(evidenced) / len(hard_found) * 100)

        # 2. Proximity, for the multi-word requirements only — a single-word skill
        #    like "Python" has nothing to be near.
        bullets = _sentences(self.experience_text)
        bullet_tokens = [_tokens(b) for b in bullets]
        multiword = [
            k['term'] for k in self.keywords
            if len(k['term'].split()) > 1 and k['type'] in ('hard', 'general')
        ]
        scattered = []
        if multiword:
            together = 0
            for term in multiword:
                words = term.split()
                if any(_within_window(t, words) for t in bullet_tokens):
                    together += 1
                elif self._found_in(self.cv_lower, term):
                    # Present on the CV, but its words never land in one bullet.
                    scattered.append(term)
            proximity_score = _clamp(together / len(multiword) * 100)
            score = _clamp(evidence_score * 0.6 + proximity_score * 0.4)
        else:
            score = evidence_score

        if scattered:
            self._recommend(
                'These phrases from the job description never appear together in a '
                'single bullet point: ' + ', '.join(f'"{s}"' for s in scattered[:4])
                + '. An ATS checks that the words of a requirement sit close '
                'together in real context, not scattered across sections.'
            )
        if evidence_score < 60:
            self._recommend(
                'Several key skills appear only in your skills list. Reference them '
                'inside your experience bullet points too.'
            )
        return score

    def check_recency(self):
        """Are the job's skills present in the two most recent roles?"""
        phase3 = self.results.get('phase3_keyword') or self.calculate_keyword_match()
        hard_found = phase3.get('hard_skills_found', [])
        if not hard_found:
            return 50

        recent_lower = self.recent_experience_text.lower()
        if not recent_lower:
            return 0
        recent = sum(1 for s in hard_found if self._found_in(recent_lower, s))
        score = _clamp(recent / len(hard_found) * 100)
        if score < 50:
            self._recommend(
                'The skills this job wants appear mainly in your older roles. '
                'Surface them in your two most recent positions, which carry the '
                'most weight.'
            )
        return score

    def check_quantification(self):
        """Do the achievement bullets carry numbers? "Increased sales 20%" ranks higher."""
        bullets = [
            b for b in _sentences(self.experience_text)
            if b.lstrip().startswith(BULLET_PREFIXES)
        ] or _sentences(self.experience_text)
        if not bullets:
            return 0

        quantified = sum(1 for b in bullets if QUANTIFIER_RE.search(b))
        score = _clamp(quantified / len(bullets) * 100)
        if score < 40:
            self._recommend(
                'Quantify your achievements. Only '
                f'{quantified} of {len(bullets)} experience bullets contain a number '
                '— "Increased organic traffic by 150%" outranks "Improved traffic".'
            )
        return score

    def run_phase4(self):
        title_score, titles = self.check_job_title_match()
        proximity = self.check_proximity()
        recency = self.check_recency()
        quantification = self.check_quantification()

        score = _clamp(
            title_score * 0.30 + proximity * 0.30 + recency * 0.20
            + quantification * 0.20
        )
        result = {
            'score': score,
            'job_title_match': title_score,
            'cv_job_titles': titles,
            'proximity_score': proximity,
            'recency_score': recency,
            'quantification_score': quantification,
        }
        self.results['phase4_context'] = result
        return result

    # -- Phase 5: experience & chronology -----------------------------------

    def check_chronology(self):
        ranges = extract_date_ranges(self.work_history_text)
        if not ranges:
            self._recommend(
                'Add explicit start and end dates to every role (e.g. "Mar 2021 - '
                'Jun 2023"). An ATS cannot compute your experience without them.'
            )
            result = {
                'score': 0, 'reverse_chronological': False, 'gaps': [],
                'job_hopping': False, 'avg_tenure_years': 0.0, 'roles_found': 0,
                'total_experience_years': 0.0,
            }
            self.results['phase5_experience'] = result
            return result

        # Reverse-chronological: each role should start no later than the one above.
        starts = [start for start, _ in ranges]
        reverse_chronological = all(
            starts[i] >= starts[i + 1] for i in range(len(starts) - 1)
        )

        # Gaps: sort oldest-first, then look for holes over 3 months.
        ordered = sorted(ranges)
        gaps = []
        for (_, prev_end), (next_start, _) in zip(ordered, ordered[1:]):
            gap_months = _months_between(prev_end, next_start)
            if gap_months > 3:
                gaps.append({
                    'from': f'{prev_end[1]:02d}/{prev_end[0]}',
                    'to': f'{next_start[1]:02d}/{next_start[0]}',
                    'months': gap_months,
                })

        tenures = [_months_between(s, e) / 12.0 for s, e in ranges]
        avg_tenure = round(sum(tenures) / len(tenures), 1) if tenures else 0.0
        short_stints = sum(1 for t in tenures if t < 1.0)
        job_hopping = short_stints >= 2 and short_stints >= len(tenures) / 2

        score = 100
        if not reverse_chronological:
            score -= 30
            self._recommend(
                'Reorder your roles so the most recent comes first. An ATS assumes '
                'reverse-chronological order and mis-reads anything else.'
            )
        if gaps:
            score -= min(20, 7 * len(gaps))
            self._recommend(
                f'Explain the {len(gaps)} employment gap(s) over 3 months '
                '(study, contract work, caring responsibilities) rather than '
                'leaving them unaccounted for.'
            )
        if job_hopping:
            score -= 20
            self._recommend(
                'Several roles are under a year. Where they were contracts or '
                'fixed-term, label them as such so they do not read as job-hopping.'
            )

        result = {
            'score': _clamp(score),
            'reverse_chronological': reverse_chronological,
            'gaps': gaps,
            'job_hopping': job_hopping,
            'avg_tenure_years': avg_tenure,
            'roles_found': len(ranges),
            'total_experience_years': total_experience_years(self.work_history_text),
        }
        self.results['phase5_experience'] = result
        return result

    # -- Phase 6: education --------------------------------------------------

    def check_education_details(self):
        education_text = self._section_text('education')
        level, label = self._highest_degree()

        gpa = self._cv_gpa()

        years = [int(y) for y in re.findall(r'\b((?:19|20)\d{2})\b', education_text or '')]
        graduation_year = max(years) if years else None

        score = 0
        if not education_text:
            self._recommend(
                'Add an EDUCATION section with your qualification, institution and '
                'dates. Most ATS platforms treat a missing education section as a '
                'parsing failure.'
            )
        else:
            score += 40
            if level:
                score += 30
            if graduation_year:
                score += 20
            else:
                self._recommend('Add your graduation year to the education section.')
            if gpa is not None:
                score += 10

        result = {
            'score': _clamp(score),
            'degree_hierarchy': label or None,
            'degree_level': level,
            'gpa': gpa,
            'graduation_year': graduation_year,
        }
        self.results['phase6_education'] = result
        return result

    # -- Phase 7: final scoring ---------------------------------------------

    def calculate_overall_score(self):
        """Weighted average of the scored phases (0-100)."""
        phase_scores = {
            'phase1_parsing': self.results['phase1_parsing']['score'],
            'phase3_keyword': self.results['phase3_keyword']['score'],
            'phase4_context': self.results['phase4_context']['score'],
            'phase5_experience': self.results['phase5_experience']['score'],
            'phase6_education': self.results['phase6_education']['score'],
        }
        total = sum(
            phase_scores[name] * weight for name, weight in PHASE_WEIGHTS.items()
        )
        return _clamp(total)

    def _sectional_scores(self):
        """Skills / experience / education breakdown, as a recruiter would read it."""
        phase3 = self.results['phase3_keyword']
        phase4 = self.results['phase4_context']
        phase5 = self.results['phase5_experience']
        phase6 = self.results['phase6_education']
        return {
            'skills': _clamp(phase3['score'] * 0.7 + phase4['proximity_score'] * 0.3),
            'experience': _clamp(
                phase5['score'] * 0.4
                + phase4['recency_score'] * 0.2
                + phase4['job_title_match'] * 0.2
                + phase4['quantification_score'] * 0.2
            ),
            'education': phase6['score'],
        }

    def text_hash(self):
        """Stable hash of the CV's content, for the deduplication check."""
        normalized = re.sub(r'\s+', ' ', self.cv_text.lower()).strip()
        return hashlib.sha256(normalized.encode('utf-8')).hexdigest()

    def get_detailed_report(self):
        """Run every phase and return the full ATS report."""
        self.results = {}
        self._recommendations = []

        self.run_phase1()
        self.run_phase2()
        self.calculate_keyword_match()
        self.run_phase4()
        self.check_chronology()
        self.check_education_details()

        overall = self.calculate_overall_score()
        threshold = getattr(settings, 'ATS_THRESHOLD', DEFAULT_THRESHOLD)

        phase1_ok = self.results['phase1_parsing']['pass']
        phase2_ok = self.results['phase2_knockout']['pass']
        rejected = not (phase1_ok and phase2_ok)

        return {
            'overall_score': overall,
            'ats_score': overall,
            'pass': (not rejected) and overall >= threshold,
            'rejected': rejected,
            'rejection_reasons': self._rejection_reasons(),
            # Split out, because the two are acted on differently: a parsing
            # failure is a problem with the CV for every job, a knock-out is
            # specific to this one.
            'parsing_failures': self._parsing_reasons(),
            'knockout_reasons': self._knockout_reasons(),
            'threshold': threshold,
            'phases': self.results,
            'sectional_scores': self._sectional_scores(),
            'recommendations': self._recommendations,
            'job_requirements': self.job_requirements,
            'text_hash': self.text_hash(),
        }

    def _rejection_reasons(self):
        """Every reason the CV would never reach a human, if it wouldn't."""
        return self._parsing_reasons() + self._knockout_reasons()

    def _parsing_reasons(self):
        """Phase 1 failures — problems with the CV itself, for any job."""
        reasons = []
        phase1 = self.results.get('phase1_parsing', {})
        if not phase1.get('pass', True):
            if phase1.get('issue'):
                reasons.append(phase1['issue'])
            if phase1.get('missing_headers'):
                reasons.append(
                    'Missing required section(s): ' + ', '.join(phase1['missing_headers'])
                )
            if phase1.get('prohibited_elements'):
                reasons.append(
                    'Contains ' + ', '.join(phase1['prohibited_elements'])
                    + ' that an ATS cannot parse'
                )
            if phase1.get('layout') and phase1['layout'] != 'single-column' \
                    and phase1['layout'] != 'unknown':
                reasons.append(f'Layout is {phase1["layout"]}, not single-column')
        return reasons

    def _knockout_reasons(self):
        """Phase 2 failures — reasons specific to this job's hard requirements."""
        reasons = []
        phase2 = self.results.get('phase2_knockout', {})
        for name in phase2.get('failed_filters', []):
            check = phase2[name]
            if name == 'experience_years':
                reasons.append(
                    f'Requires {check["required"]} years of experience; CV shows '
                    f'{check["found"]}'
                )
            elif name == 'education':
                reasons.append(
                    f'Requires a {check["required"]} degree; CV shows '
                    f'{check["found"] or "none"}'
                )
            elif name == 'certifications':
                reasons.append(
                    'Missing required certification(s): '
                    + ', '.join(c.upper() for c in check['missing'])
                )
            elif name == 'gpa':
                reasons.append(
                    f'Requires a minimum GPA of {check["required"]}; CV shows '
                    f'{check["found"]}'
                )
            elif name == 'location':
                reasons.append(
                    f'Role is in {check["required"]}; no matching location or '
                    f'relocation statement on the CV'
                )
            elif name == 'work_authorization':
                reasons.append('CV requires visa sponsorship; this role does not offer it')
        return reasons


# ---------------------------------------------------------------------------
# Convenience entry points
# ---------------------------------------------------------------------------

def check_cv_against_job(cv_text, job_description, job_title='', job_location='',
                         file_path=None):
    """Run the full ATS check and return the report dict.

    The convenience wrapper used across the app: mines the job's requirements
    from its description, then runs all seven phases.
    """
    requirements = extract_job_requirements(job_description, job_title, job_location)
    checker = ATSChecker(cv_text, job_description, requirements, file_path=file_path)
    return checker.get_detailed_report()


def check_cv_format(cv_text, file_path=None):
    """Phase 1 only — is this CV ATS-readable at all, regardless of any job?

    Used at upload time, where there is no job description to score against.
    """
    checker = ATSChecker(cv_text, '', job_requirements={}, file_path=file_path)
    return checker.run_phase1() | {'recommendations': checker._recommendations}
